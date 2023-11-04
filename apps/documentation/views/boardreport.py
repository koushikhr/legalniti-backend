from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.views import View
import requests
import json
import urllib.parse
from django.http import HttpResponseBadRequest
import io
import urllib.request
import tempfile
from PIL import Image
from easyocr import Reader
from ultralytics.models.yolo import YOLO
from django.shortcuts import render
from django.http import HttpResponse
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document
from rest_framework.views import APIView
from rest_framework.response import Response
from django.views.generic import TemplateView
from docx import Document
from docx.shared import Pt
import openpyxl
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import requests
from rest_framework import status
import boto3
from botocore.exceptions import NoCredentialsError
from io import BytesIO
from datetime import datetime
import pytz
import datetime
from docx2pdf import convert
from docx.shared import Inches
import os
import re


class BOARDREPORT(APIView):

    def post(self,request,*args,**kwargs):


        company_name = request.data.get('companyname')
        end_of_fy = request.data.get('end_of_fy')
        start_of_fy = request.data.get('start_of_fy')
         #website of the company
        company_web=''
        num_of_boardmeeting= int(request.data.get('num_of_meet',0))
        directors_count= int(request.data.get('num_of_directors',0))
        company_data=request.data.get('company_activity',{})
        dividend_data_len=int(request.data.get('selectedData',0))
        shareholders_count=int(request.data.get('shareHolders',0))



        current_datetime = datetime.datetime.now()

        # Format the current date and time
        formatted_time = current_datetime.strftime("%Y-%m-%d_%H-%M-%S")

        doc = Document()

        # Add a title above the heading
        title = doc.add_paragraph('On the letterhead of the Company')
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add a heading to the document
        heading = doc.add_paragraph('DIRECTORS’ REPORT')
        heading.style = 'Heading1'
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # Function to add a heading with specified text
        def add_heading(doc, text, level=1):
            heading = doc.add_paragraph(text, style=f'Heading{level}')
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        # Function to add content paragraph with specified text
        def add_content(doc, text):
            content = doc.add_paragraph(text)
            content.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            
        # Define a function to add a table
        def add_table(doc, heading, data):
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            
            # Set table header row
            header_row = table.rows[0]
            header_cells = header_row.cells
            for i, column_heading in enumerate(heading):
                cell = header_cells[i]
                cell.text = column_heading
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                cell.paragraphs[0].bold = True
                cell.paragraphs[0].font.size = Pt(10)
            
            # Populate table with data
            for row_data in data:
                row = table.add_row().cells
                for i, value in enumerate(row_data):
                    row[i].text = value
                    row[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    row[i].paragraphs[0].font.size = Pt(10)

        
        content = doc.add_paragraph()
        content.add_run('To,\n')
        content.add_run('The Members,\n')
        content.add_run('Company Name')
        content.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        # Add the introduction sentence
        intro = doc.add_paragraph()
        intro.add_run(f'Your directors have immense pleasure in presenting the Board Report of {company_name} for the financial year ended {end_of_fy}.')
        intro.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Add the period of report heading
        period_heading = doc.add_paragraph('1) PERIOD OF REPORT:')
        period_heading.style = 'Heading2'
        period_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        content = doc.add_paragraph()
        content.add_run(f'This report pertains to the period from {start_of_fy} to {end_of_fy}')


        period_heading = doc.add_paragraph('2)WEBLINK OF ANNUAL RETURN: (if the company has website)')
        period_heading.style = 'Heading2'
        period_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        weblink_content = doc.add_paragraph()
        weblink_content.add_run(f'The Company is having a website, i.e., {company_web}, and the annual return of the Company has been published on such website. The link of the same is given below:\n')
        weblink_content.add_run(" OR\n" )
        weblink_content.add_run("The Company doesn’t have any website.  Therefore, no need to of publication of Annual Return.")
        weblink_content.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Add a heading for the financial summary
        financial_summary_heading = doc.add_paragraph('3) FINANCIAL SUMMARY:')
        financial_summary_heading.style = 'Heading2'
        financial_summary_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Add content to the financial summary section
        financial_summary_content = doc.add_paragraph()
        financial_summary_content.add_run("The following gives a summary of the financial results of the Company:")
        financial_summary_content.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Add a table with the specified headings and columns
        table = doc.add_table(rows=25, cols=3)
        table.style = 'Table Grid'  # Apply a table style

        # Populate the header row of the table
        heading_cells = table.rows[0].cells
        heading_cells[0].text = "Particulars"
        heading_cells[1].text = "Financial Year relating to the current reporting period ended 31.03.2021"
        heading_cells[2].text = "Financial Year relating to the previous reporting period ended 31.03.2020"

        for i, financial_data in enumerate(request.data.get('selectedValues', [])):
            
            particular = financial_data.get('Particulars', '')
            current_period = financial_data.get('current','')
            previous_period = financial_data.get('previous','')

            row_cells = table.rows[i].cells
            row_cells[0].text = particular
            row_cells[1].text = current_period
            row_cells[2].text = previous_period
        
        # Adjust column widths
        col_widths = [Inches(2.5), Inches(2.5), Inches(2.5)]
        for i, width in enumerate(col_widths):
            table.columns[i].width = width

        # Add a heading for Consolidated Financial Statements
        consolidated_heading = doc.add_paragraph('4) CONSOLIDATED FINANCIAL STATEMENTS:', style='Heading2')
        consolidated_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Add the content for Consolidated Financial Statements
        consolidated_content = doc.add_paragraph(
            "Company doesn’t have any subsidiaries. Hence, preparation of consolidated financial statement for the financial year 2020-21 is not applicable for the Company.\n"
            "OR\n"
            "If company has subsidiary: "
        )
        consolidated_content.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        meeting_text = (
            f"{num_of_boardmeeting} Board meetings held during the Financial Year ended 31st March, 2021 under review on the following dates. The intervening gap between the Meetings was within the period prescribed under the Companies Act, 2013. "
        )
        meeting_paragraph = doc.add_paragraph(meeting_text)
        meeting_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


        # Read the JSON data for Board meetings
        

        # Add a table for Board meetings
        board_meeting_table = doc.add_table(rows=num_of_boardmeeting+ 1, cols=5)
        board_meeting_table.style = 'Table Grid'  # Apply a table style

        # Populate the header row of the board meeting table
        board_meeting_heading_cells = board_meeting_table.rows[0].cells
        board_meeting_heading_cells[0].text = "S. No."
        board_meeting_heading_cells[1].text = "Date of Meeting"
        board_meeting_heading_cells[2].text = "No. of gap in days between two meetings"
        board_meeting_heading_cells[3].text = "Total No. of Directors entitled to attend the meeting"
        board_meeting_heading_cells[4].text = "Total No. of Directors attended the Meeting"

        for i, board_meet_data in enumerate(request.data.get('selectedBoard', [])):
            sr_no = str(i + 1)
            date = board_meet_data.get('date', '')
            gap_days = board_meet_data.get('gap_days','')
            entitled = board_meet_data.get('directors_entitled','')
            attended = board_meet_data.get('directors_attended','')


            row_cells = table.rows[i].cells
            row_cells[0].text = sr_no
            row_cells[1].text = date
            row_cells[2].text = gap_days
            row_cells[3].text = entitled
            row_cells[4].text = attended
        
        # Adjust column widths for the board meeting table
        board_meeting_col_widths = [Inches(0.7), Inches(1.5), Inches(1.5), Inches(1.8), Inches(1.8)]
        for i, width in enumerate(board_meeting_col_widths):
            board_meeting_table.columns[i].width = width

        # Add a paragraph for the names of Directors and their attendance
        attendance_paragraph = doc.add_paragraph()
        attendance_paragraph.add_run("The names of Directors on the Board and their attendance at the Board Meetings are as under:")
        attendance_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Add a table for Directors attendance
        directors_attendance_table = doc.add_table(rows= directors_count+ 1, cols=3)
        directors_attendance_table.style = 'Table Grid'  # Apply a table style

        # Populate the header row of the Directors attendance table
        attendance_heading_cells = directors_attendance_table.rows[0].cells
        attendance_heading_cells[0].text = "Name of Directors"
        attendance_heading_cells[1].text = "Total Number of Meetings held during the Financial Year"
        attendance_heading_cells[2].text = "Number of Meetings attended"

        for i, director_data in enumerate(request.data.get('selectedDirectors', [])):
                
            
            name = director_data.get('NAME', '')
            meeting_held=director_data.get('meeting_held','')
            meeting_attended = director_data.get('meeting_attended','')

            row_cells = directors_attendance_table.rows[i].cells
            row_cells[0].text = name
            row_cells[1].text = meeting_held
            row_cells[2].text = meeting_attended

        attendance_col_widths = [Inches(2.5), Inches(2.5), Inches(2.5)]
        for i, width in enumerate(attendance_col_widths):
                directors_attendance_table.columns[i].width = width

            # Add a heading for Director's Responsibility Statement
        responsibility_heading = doc.add_paragraph('5) DIRECTOR’S RESPONSIBILITY STATEMENT:', style='Heading2')
        responsibility_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            # Add your content for the Director's Responsibility Statement here
        responsibility_statement = (
                "Pursuant to Section 134(3)(c) and 134(5) of the Companies Act, 2013 the Board of Directors of the Company confirms that-\n\n"
                "(a) in the preparation of the annual accounts, the applicable accounting standards had been followed along with proper explanation relating to material departures;\n\n"
                "(b) the directors had selected such accounting policies and applied them consistently and made judgments and estimates that are reasonable and prudent so as to give a true and fair view of the state of affairs of the company at the end of the financial year and of the profit and loss of the company for that period;\n\n"
                "(c) the directors had taken proper and sufficient care for the maintenance of adequate accounting records in accordance with the provisions of this Act for safeguarding the assets of the company and for preventing and detecting fraud and other irregularities;\n\n"
                "(d) the directors had prepared the annual accounts on a going concern basis; and\n\n"
                "(e) The Company being unlisted, sub clause (e) of section 134(5) of the Companies Act, 2013 pertaining to laying down internal financial controls to be followed by the company and that such internal financial controls are adequate and were operating effectively is not applicable to the Company.\n\n"
                "(f) the directors had devised proper systems to ensure compliance with the provisions of all applicable laws and that such systems were adequate and operating effectively."
            )

        # Add the Director's Responsibility Statement content
        responsibility_paragraph = doc.add_paragraph(responsibility_statement)
        responsibility_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT



        # Add a heading for Details of Fraud reported by the Auditor
        fraud_heading = doc.add_paragraph('7) DETAILS OF FRAUD REPORTED BY THE AUDITOR:', style='Heading2')
        fraud_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Add your content for Details of Fraud reported by the Auditor here
        fraud_details = (
            "As per auditors’ report, no fraud u/s 143(12) reported by the auditor."
        )

        # Add the Details of Fraud reported by the Auditor content
        fraud_paragraph = doc.add_paragraph(fraud_details)
        fraud_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


        # Add a heading for Explanations or Comments by the Board
        explanations_heading = doc.add_paragraph('8) EXPLANATIONS OR COMMENTS BY THE BOARD ON EVERY QUALIFICATION, RESERVATION OR ADVERSE REMARK OR DISCLAIMER MADE:', style='Heading2')
        explanations_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Add your content for Explanations or Comments by the Board
        explanations_content = (
            "Statutory Auditors:\n"
            "The Statutory Auditors’ Report for the financial year ended March 31, 2021 does not contain any qualification, reservation or adverse remark.\n\n"
            "Reply by the board:\n"
            "The observations of the Statutory Auditors, when read together with the relevant notes to the accounts and accounting policies have clear opinion on the Financial Statements for year ended March 31, 2021. There has been due compliance of all statutory rules, regulations and standards.\n\n"
            "Secretarial Audit:\n"
            "Secretarial Audit is not applicable to our company. (Mention, if applicable)"
        )

        # Add the Explanations or Comments by the Board content
        explanations_paragraph = doc.add_paragraph(explanations_content, style='BodyText')
        explanations_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Add a heading for Particulars of Loans, Guarantees or Investments under Section 186
        section_186_heading = doc.add_paragraph('9) PARTICULARS OF LOANS, GUARANTEES OR INVESTMENTS MADE UNDER SECTION 186:', style='Heading2')
        section_186_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


        # Add the content for Particulars of Loans, Guarantees or Investments under Section 186
        section_186_content = doc.add_paragraph(
            "If the company has not given any Loan, Guarantee or any investment u/s 186:\n"
            "The Company has not given any Loan, Guarantee or any investment in the Financial Year under section 186 of the Companies Act of 2013.\n\n"
            "If given,\n"
            "Particulars of the loans given, investments made, guarantees given or securities provided during the year and the purpose for which the loans / guarantees /securities are proposed to be utilized by the recipient of such loan / guarantee/ security shall be reported."
        )

        section_186_content.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Add a heading for Particulars of Contracts under Section 188
        section_188_heading = doc.add_paragraph('10) PARTICULARS OF CONTRACTS UNDER SECTION 188:', style='Heading2')
        section_188_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Add the content for Particulars of Contracts under Section 188
        section_188_content = doc.add_paragraph(
            "If no Contracts are there:\n"
            "All related party transactions that were entered during the financial year/period ended 31st March 2021, were on an arm’s length basis and were in the ordinary course of business. Therefore, the provisions of Section 188 of the Companies Act, 2013 were not attracted. \n"
            "Further, there are no materially significant related party transactions during the year under review made by the Company with Promoters, Directors, or other designated persons which may have a potential conflict with the interest of the Company at large. Thus, disclosure in Form AOC-2 is not required. However, the disclosure of transactions with related party for the year, as per Accounting Standard -18 Related Party Disclosures is given in Note no __ _________ to the Balance Sheet as on 31st March 2021.\n\n"
            "OR\n\n"
            "If Contracts are there:\n"
            "The Company has entered various Related Parties Transactions as defined under Section 188 of the Companies Act, 2013 with related parties as defined under Section 2 (76) of the said Act. Further all the necessary details relating to related party transactions during the financial year/period under review are provided in AOC-2 as “Annexure ..”."
        )
        section_188_content.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Add a heading for State of Company's Affairs
        affairs_heading = doc.add_paragraph('11) STATE OF COMPANY’S AFFAIRS:', style='Heading2')
        affairs_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        affairs_content = (
            f"The Company is engaged in the business {company_data.business_activity} and {company_data.business_activity_2}. "
            f"There has been no change in the business of the Company during the financial year ended March 31, 2021.\n\n"
            "The highlights of the Company’s performance are as under:-\n\n"
            f"Net Profit for the year increased from Rs. {company_data.net_profit_previous_year} to Rs. {company_data.net_profit_current_year}\n\n"
            f"Earnings per share have increased from Rs. {company_data.earnings_per_share_previous} to Rs. {company_data.earnings_per_share_current}\n\n"
            f"During the year company has earned a net profit of Rs. {company_data.net_profit_current_year}. "
            "Your Directors are confident that going forward, the Company shall perform well."
        )

        # Add the State of Company's Affairs content
        affairs_paragraph = doc.add_paragraph(affairs_content)
        affairs_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        # Add a heading for Impact of COVID-19
        covid_impact_heading = doc.add_paragraph('12) IMPACT OF COVID-19:', style='Heading2')
        covid_impact_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Add a heading for Amount to Transfer to Reserves
        reserves_heading = doc.add_paragraph('13) AMOUNT IF ANY, WHICH THE BOARD PROPOSES TO TRANSFER TO ANY RESERVES:', style='Heading2')
        reserves_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Add the content for Amount to Transfer to Reserves
        reserves_content = doc.add_paragraph(
            "If amount is proposed to be transferred:\n"
            "The Company proposes to transfer a sum of Rs_________ to ___________ Reserve during the financial year ended 31st March 2021.\n\n"
            "OR\n\n"
            "If amount is not proposed to be transferred:\n"
            "The Board of Directors of your company has decided not to transfer any amount to the Reserves for the year under review."
        )
        reserves_content.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Add a heading for Transfer of Amounts to Investor Education and Protection Fund
        iepf_heading = doc.add_paragraph('14) TRANSFER OF AMOUNTS TO INVESTOR EDUCATION AND PROTECTION FUND:', style='Heading2')
        iepf_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        combined_data_content = (
            "If company is not recommending dividend:\n"
            "The Board of Directors of your company, after considering holistically the relevant circumstances has decided that it would be prudent, not to recommend any Dividend for the year under review.\n\n"
            "If board was recommended any dividend during the financial year:\n"
            "The Board has recommended interim dividend/Final Dividend paid to the ordinary equity/Preference shareholders of the Company as per below table out of surplus in profit and loss account as appearing in the signed audited financials as at 31st March 2020 and pro-forma profit details for the year till date 31st December 2020."
        )

        # Add the "Combined Data" content as a paragraph
        combined_data_paragraph = doc.add_paragraph(combined_data_content)
        combined_data_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        dividend_table = doc.add_table(rows=len(dividend_data_len) + 1, cols=5)
        dividend_table.style = 'Table Grid'  # Apply a table style

        dividend_heading_cells = dividend_table.rows[0].cells
        dividend_heading_cells[0].text = "Sr. No."
        dividend_heading_cells[1].text = "Declaration Date"
        dividend_heading_cells[2].text = "Dividend per Equity Shares"
        dividend_heading_cells[3].text = "Total Interim Dividend Paid (in Rs.)"
        dividend_heading_cells[4].text = "Total Tax paid (in Rs.)"

        for i, dividend_data in enumerate(request.data.get('selectedDividend', [])):
                
            sr_no = str(i + 1)
            declaration_date = dividend_data.get('declaration_date', '')
            dividend_per_equity=dividend_data.get('dividend_per_equity','')
            total_interim_dividend = dividend_data.get('total_interim_dividend','')
            total_tax_paid=dividend_data.get('total_tax_paid','')
            total_overflow=dividend_data.get('total_overflow','')


            row_cells = dividend_table.rows[i].cells
            row_cells[0].text = sr_no
            row_cells[1].text = declaration_date
            row_cells[2].text = dividend_per_equity
            row_cells[3].text = total_interim_dividend
            row_cells[4].text = total_tax_paid
        
        # Adjust column widths for the Dividend Details table
        dividend_col_widths = [Inches(0.7), Inches(1.5), Inches(1.5), Inches(1.8), Inches(1.8)]
        for i, width in enumerate(dividend_col_widths):
            dividend_table.columns[i].width = width

        # Add a table for Shareholders Details
        shareholders_table = doc.add_table(rows=shareholders_count + 1, cols=6)
        shareholders_table.style = 'Table Grid'  # Apply a table style

        # Populate the header row of the Shareholders Details table
        shareholders_heading_cells = shareholders_table.rows[0].cells
        shareholders_heading_cells[0].text = "Sr. No."
        shareholders_heading_cells[1].text = "Shareholders Name"
        shareholders_heading_cells[2].text = "Number of Shares held"
        shareholders_heading_cells[3].text = "% Holding"
        shareholders_heading_cells[4].text = "Face value Per Share (INR)"
        shareholders_heading_cells[5].text = "Paid-up Capital (INR)"

        for i, shareholder_data in enumerate(request.data.get('shareholderData', [])):
                
            sr_no = str(i + 1)
            name = shareholder_data.get('shareholder_name', '')
            number_of_shares=shareholder_data.get('num_of_share','')
            percentage_holding = shareholder_data.get('percentage_holding','')
            face_value_per_share=shareholder_data.get('face_value_per_share','')
            paid_up_capital=shareholder_data.get('paid_up_capital','')


            row_cells = dividend_table.rows[i].cells
            row_cells[0].text = sr_no
            row_cells[1].text = name
            row_cells[2].text = number_of_shares
            row_cells[3].text = percentage_holding
            row_cells[4].text = face_value_per_share
            row_cells[5].text = paid_up_capital

        # Adjust column widths for the Shareholders Details table
        shareholders_col_widths = [Inches(0.7), Inches(2.5), Inches(1.5), Inches(1.5), Inches(1.5), Inches(1.5)]
        for i, width in enumerate(shareholders_col_widths):
            shareholders_table.columns[i].width = width
        dividend_content = doc.add_paragraph(
            "No Profit:\n"
            "The Board of Directors of your company has not declared any Dividend for the current financial year due to conservation of Profits/due to loss incurred by the Company/due to insufficient profit."
        )
        dividend_content.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Add a heading for Major Events Occurred During the Year
        events_heading = doc.add_paragraph('16) MAJOR EVENTS OCCURRED DURING THE YEAR:', style='Heading2')
        events_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
        # Add the content for Major Events Occurred During the Year
        # Add the content for Major Events Occurred During the Year
        events_content = (
            "Change in the nature of business/ Status of the Company:\n"
            "There have been no changes made in the nature of the business by your Company for the year to which the financial statements and the report relate to.\n\n"
            "Change in the financial year:\n"
            "There have been no changes made in the financial year by your Company for the year to which the financial statements and the report relate to.\n\n"
            "Details and status of acquisition, merger, expansion, modernization and diversification:\n"
            "There have been no acquisition, merger, expansion, modernization and diversification by your Company for the year to which the financial statements and the report relate to.\n\n"
            "Developments, acquisition and assignment of material Intellectual Property Rights:\n"
            "There have been no developments, acquisition and assignment of material Intellectual Property Rights by your Company for the year to which the financial statements and the report relate to.\n\n"
            "Material changes and commitments, if any, affecting the financial position of the company, having occurred since the end of the year and till the date of the Report:\n\n"
            "If no change took place:\n"
            "There have been no material changes and commitments, which affect the financial position of the Company which have occurred between the end of the financial year to which the financial statements relate and the date of this Report.\n\n"
            "If any change took place:\n"
            "There have been the following material changes and commitments which could affect the financial position of the Company, which have occurred between the end of the financial year of the company to which the financial statements relate and the date of the report:\n"
            "1.\n"
            "2."
        )

        # Add the Major Events Occurred During the Year content
        events_paragraph = doc.add_paragraph(events_content)
        events_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Add a heading for CONSERVATION OF ENERGY, TECHNOLOGY ABSORPTION, FOREIGN EXCHANGE EARNINGS/OUTGO
        conservation_heading = doc.add_paragraph('17)CONSERVATION OF ENERGY, TECHNOLOGY ABSORPTION, FOREIGN EXCHANGE EARNINGS/OUTGO:', style='Heading2')
        conservation_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Add a table for Foreign Exchange Earnings and Expenditures
        fx_table = doc.add_table(rows=4, cols=3)
        fx_table.style = 'Table Grid'

        # Set column widths
        column_widths = [4.0, 4.0, 4.0]
        for i, width in enumerate(column_widths):
            fx_table.columns[i].width = Inches(width)

        # Add table headers
        headers = fx_table.rows[0].cells
        headers[0].text = "Particulars"
        headers[1].text = "Financial Year ending for the current reporting period"
        headers[2].text = "Financial Year ending for the previous reporting period"

        # Add table data
        data_rows = [
            ("Foreign Exchange Earnings", "", ""),
            ("Foreign Exchange", "", ""),
            (" Expenditures", "", "")
        ]

        for i, row in enumerate(fx_table.rows[1:]):
            for j, value in enumerate(data_rows[i]):
                row.cells[j].text = value

        # Add a heading for Risk Management
        risk_heading = doc.add_paragraph('18)RISK MANAGEMENT:', style='Heading2')
        risk_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        # Add content using the add_content function
        add_content(doc, "The Company does not have any Risk Management Policy as the elements of risk threatening the Company’s existence are very minimal.")

        # OR

        add_content(doc, "A statement indicating the development and implementation of a risk management policy for the company. Such statement shall, inter alia, disclose:")
        add_content(doc, "(a) various elements of risk which, in the opinion of the Board, may threaten the existence of the company and")
        add_content(doc, "(b) strategy to mitigate such risks.")

        # Add a heading for Corporate Social Responsibility (CSR)
        csr_heading = doc.add_paragraph('CORPORATE SOCIAL RESPONSIBILITY (CSR):', style='Heading2')
        csr_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Add the content for Corporate Social Responsibility (CSR)
        csr_content = (
            "The requirement of constitution of CSR Committee is not applicable to the Company.\n\n"
            "OR\n\n"
            "Pursuant to the provisions of Section 135 of the Companies Act, 2013 read with the Companies (Corporate Social Responsibility Policy) Rules, 2014, your Company as part of its CSR initiatives has undertaken projects/programs in accordance with the CSR Policy.\n\n"
            "The Members of CSR Committee being as follows:\n\n"
            "1.\n"
            "2.\n"
            "3.\n"
            "4.\n\n"
            "The Company is required to spend a gross amount of Rs. __________/- for the year ending March 31, 2021. However, the Company has spent more than the required amount and in total the amount spent on CSR activity during the year is Rs. ___________/-. The Company has met the statutory limit to be spent for CSR Expenditure for FY 2018-19, FY 2019-20 and FY 2020-21. As on March 31, 2021, there is no unspent CSR amount in the books of Company.\n\n"
            "The annual report on CSR Activities is enclosed as per prescribed format as Annexure __ and forms part of this report."
        )

        # Add the CSR content
        csr_paragraph = doc.add_paragraph(csr_content)
        csr_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Add a heading for Statutory Auditors
        auditors_heading = doc.add_paragraph('STATUTORY AUDITORS:', style='Heading2')
        auditors_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Add content for Statutory Auditors
        auditors_content = doc.add_paragraph(
            "Names of the Statutory Auditor, Cost Auditor and Secretarial Auditor and details of any change in such Auditors, during the year and up to the date of the Report due to resignation / casual vacancy / removal / completion of term shall be disclosed in the Report."
        )
                
        auditors_content.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        # Add a heading for Compliance with Secretarial Standards
        secretarial_heading = doc.add_paragraph('COMPLIANCE WITH SECRETARIAL STANDARDS:', style='Heading2')
        secretarial_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


        # Add the content for Compliance with Secretarial Standards
        secretarial_content = doc.add_paragraph("The Company complies with all applicable secretarial standards issued by the Institute of Company Secretaries of India.")
        secretarial_content.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        # Add a heading for Details of Directors or Key Managerial Personnel Appointments/Resignations
        appointments_heading = doc.add_paragraph('THE DETAILS OF DIRECTORS OR KEY MANAGERIAL PERSONNEL WHO WERE APPOINTED OR HAVE RESIGNED DURING THE YEAR:', style='Heading2')
        appointments_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Add a table for changes in Board constitution
        appointments_table = doc.add_table(rows=3, cols=6)
        appointments_table.style = 'Table Grid'

        # Add table headers
        table_headers = appointments_table.rows[0].cells
        table_headers[0].text = 'Sl No.'
        table_headers[1].text = 'Name'
        table_headers[2].text = 'Particulars'
        table_headers[3].text = 'Designation'
        table_headers[4].text = 'Date of Appointment'
        table_headers[5].text = 'Date of Cessation'

        # Add table data
        data = [
            (1, 'John Doe', 'Appointed as Director', 'Director', '2023-01-15', '2023-07-30'),
            (2, 'Jane Smith', 'Resigned as CFO', 'CFO', '2022-05-20', '2023-06-15')
            # Add more rows as needed
        ]

        for row_data in data:
            row = appointments_table.add_row().cells
            for i, value in enumerate(row_data):
                row[i].text = str(value)

                
        # Add content about appointment of KMP
        kmp_content = doc.add_paragraph(
            "In view of the applicable provisions of the Companies Act, 2013, the Company is not mandatorily required to appoint any whole time Key Managerial Personnel (KMP)."
        )
        kmp_content.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


        # Add heading and content about Company's policies
        policies_heading = doc.add_paragraph('COMPANY’S POLICIES ON APPOINTMENT OF DIRECTORS, REMUNERATION AND OTHER MATTERS:', style='Heading2')
        policies_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Add content about Section 178
        section_178_content = doc.add_paragraph(
            "The Company doesn’t fall under the purview of the criteria laid in Section 178 of the Companies Act, 2013 read with rule 6 of meeting of Board & its powers rules, 2014. Therefore, reporting under this head shall not apply to the Company."
        )
        section_178_content.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Add a heading for Managerial Remuneration
        remuneration_heading = doc.add_paragraph('MANAGERIAL REMUNERATION:', style='Heading2')
        remuneration_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Add content about particulars under section 134
        particulars_content = doc.add_paragraph(
            "Particulars in terms of section 134 of the Companies Act, 2013, read with Companies (Appointment and Remuneration of Managerial Personnel) Rules, 2014. There was increase or decrease in the remuneration of any Executive Director, Chief Financial Officer or Company Secretary of the Company. Following are the details of the same:"
        )
        particulars_content.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


        # Add a table for Managerial Remuneration
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        table.allow_autofit = False

        # Add table headings
        row = table.rows[0]
        row.cells[0].text = 'Name'
        row.cells[1].text = 'Designation'
        row.cells[2].text = 'Remuneration in FY 2020-21'
        row.cells[3].text = 'Remuneration in FY 2019-20'
        row.cells[4].text = 'Increase or Decrease during FY 2020-21'

        # Add rows to the table
        data = [
            ['Director 1', 'Executive Director', 'Rs. X', 'Rs. Y', 'Rs. Z'],
            ['Director 2', 'CFO', 'Rs. A', 'Rs. B', 'Rs. C']
        ]

        for item in data:
            row = table.add_row().cells
            for i, value in enumerate(item):
                row[i].text = value

                
        add_heading(doc, "SHARE CAPITAL AND ITS CHANGES", level=1)
        # Add the content for SHARE CAPITAL AND ITS CHANGES
        add_content(doc, [
            "During the financial year, there is no change in the capital structure of the Company. The Company has not made any allotment of equity shares to any of the shareholders.\n",
            "",
            "If any changes as mentioned below occur, the same should be disclosed:\n",
            "(a) Change in the authorized, issued, subscribed, and paid-up share capital;\n",
            "(b) Reclassification or sub-division of the authorized share capital;\n",
            "(c) Reduction of share capital or buyback of shares;\n",
            "(d) Change in the capital structure resulting from restructuring; and\n",
            "(e) Change in voting rights."
        ])

        # Add the heading for DETAILS OF SUBSIDIARY, JOINT VENTURE OR ASSOCIATE COMPANIES
        add_heading(doc, "DETAILS OF SUBSIDIARY, JOINT VENTURE OR ASSOCIATE COMPANIES:", level=1)

        # Add the content for no Subsidiary, Joint venture or Associate Company
        add_content(doc, "Your Company does not have any Subsidiary, Joint venture or Associate Company.\n\n"
                        "OR\n\n"
                        "The names of companies which have become or ceased to be its subsidiaries, joint ventures or associate companies during the year:")

        # Add the table heading
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Table Grid'
        header_cells = table.rows[0].cells
        header_cells[0].text = "S. No."
        header_cells[1].text = "Name of the company"
        header_cells[2].text = "Subsidiary/Joint Venture/Associate"
        header_cells[3].text = "No. of Shares Held"
        header_cells[4].text = "Percentage of Shareholding"
        header_cells[5].text = "Become/ceased"
        header_cells[6].text = "Effective date"

        # Add data rows to the table
        data = [
            ["1", "Company A", "Subsidiary", "1000", "50%", "Become", "01-01-2023"],
            ["2", "Company B", "Joint Venture", "500", "30%", "Ceased", "15-06-2023"]
        ]

        for row in data:
            row_cells = table.add_row().cells
            for i in range(len(row)):
                row_cells[i].text = row[i]
        heading = doc.add_heading("EMPLOYEE STOCK OPTION PLAN (ESOP)", level=1)

        # Add the content for ESOP
        esop_content = doc.add_paragraph(
            "If company has not issued:\n"
            "During the year there was no Employee Stock Option Plan in the Company.\n\n"
            "If company issued ESOP:\n"
            "The details of Employee stock options for the year _______________ are as follows:"
        )
        esop_content.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


        heading = doc.add_heading("INTERNAL FINANCIAL CONTROL AND THEIR ADEQUACY:", level=1)

        add_content(doc, "The Company has in place adequate internal financial controls with reference to financial statements. During the year, such controls were tested and no reportable material weakness in the design or operation was observed.")
        heading = doc.add_heading ("GENERAL:")

        content = (
            "Your Directors state that no disclosure or reporting is required in respect of the following items as there were no transactions on these items during the year under review:\n\n"
            "• The details of application made or any proceeding pending under the Insolvency and Bankruptcy Code, 2016 (31 of 2016) during the year alongwith their status as at the end of the financial year.\n"
            "• The details of difference between amount of the valuation done at the time of one time settlement and the valuation done while taking loan from the Banks or Financial Institutions along with the reasons thereof."
        )

        # Use the add_content function to add the content to the document
        add_content(doc, content)

        heading = doc.add_heading("DISCLOSURE UNDER THE SEXUAL HARASSMENT OF WOMEN AT WORKPLACE (PREVENTION, PROHIBITION, AND REDRESSAL) ACT, 2013:")

        # Define the content to be added
        content = (
            "The company has in place a Policy for prevention of Sexual Harassment at the Workplace in line with the requirements of the Sexual Harassment of Women at the Workplace (Prevention, Prohibition & Redressal) Act, 2013 and also has a policy and framework for employees to report sexual harassment cases at workplace and its process ensures complete anonymity and confidentiality of information. Adequate workshops and awareness programmes against sexual harassment are conducted across the organization.\n\n"
            "The following is a summary of sexual harassment complaints received and disposed of during the year under review:\n\n"
            "Number of complaints pending at the beginning of the year: __________\n"
            "Number of complaints received during the year: __________\n"
            "Number of complaints disposed of during the year: __________\n"
            "Number of cases pending at the end of the year: __________\n\n"
            "During the Financial Year under review, the Internal Complaints Committee under the Sexual Harassment of Women at Workplace (Prevention, Prohibition and Redressal) Act, 2013 was reconstituted with following members:\n"
        )

        # Use the add_content function to add the content to the document
        add_content(doc, content)

        heading = doc.add_heading("THE DETAILS RELATING TO DEPOSITS, COVERED UNDER CHAPTER V OF THE ACT:")
        content = ("The Company has not accepted any deposits during the year under review.\n"
                "OR\n"
                "(a) accepted during the year;\n"
                "(b) remained unpaid or unclaimed as at the end of the year;\n"
                "(c) whether there has been any default in repayment of deposits or payment of interest thereon during the year and if so, number of such cases and the total amount involved-\n"
                "(i) at the beginning of the year;\n(ii) maximum during the year;\n(iii) at the end of the year;"
        )
        add_content(doc, content)

        heading = doc.add_heading("THE DETAILS OF DEPOSITS WHICH ARE NOT IN COMPLIANCE WITH THE REQUIREMENTS OF CHAPTER V OF THE ACT; ")
        add_content(doc,"The Company has not accepted any deposits which are not in compliance of the Companies Acceptance of Deposits) Rules, 2014 during the year.")

        heading = doc.add_heading("EVENT SUBSESQUENT TO THE DATE OF FINANCIAL STATEMENT:")
        # Function to add content to the document
        def add_content(doc, content):
            paragraph = doc.add_paragraph(content)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        # Create a new document


        # Define the content to be added
        content = (
            "There are no material changes and commitments affecting the financial position "
            "of the company occurred between the end of the financial year to which this "
            "financial statement relate to the date of this report."
        )

        # Add the content to the document using the function
        add_content(doc, content)


        # Heading and content for "DISCLOSURE  ON  ESTABLISHMENT  OF  VIGIL MECHANISM:"
        heading = doc.add_heading("DISCLOSURE  ON  ESTABLISHMENT  OF  VIGIL MECHANISM:", level=1)
        content = ("The provisions of Section 177(9) of the Companies Act, 2013 read with Rule 7 of the Companies (Meetings of the Board and its Powers) Rules, 2013 are not applicable to the Company.")
        add_content(doc, content)

        # Heading and content for "DISCLOSURE UNDER SEC 67(3) READ WITH  RULE 16 OF COS(SHARE CAP & DEBENTURE) RULES 2014 DISCLOSURE IN RESPECT OF VOTING RIGHTS NOT EXERCISED DIRECTLY BY THE EMPLOYEES:"
        heading = doc.add_heading("DISCLOSURE UNDER SEC 67(3) READ WITH  RULE 16 OF COS(SHARE CAP & DEBENTURE) RULES 2014 DISCLOSURE IN RESPECT OF VOTING RIGHTS NOT EXERCISED DIRECTLY BY THE EMPLOYEES:", level=1)
        add_content(doc, "")

        heading = doc.add_heading("ACKNOWLEDGMENT AND APPRECIATION:")
        content = ("Your Directors would like to express their sincere appreciation for the assistance and cooperation received from the banks, Government authorities, customers, vendors and members during the year under review. Your Directors also wish to place on record their deep sense of appreciation for the committed services by the Company’s executives, staff and workers.")

        add_content = (doc , content)

        # Define the content to be added
        content = (
            "For and on behalf of the Board of,\n"
            "______________________________\n\n"
            "_______________\t\t___________________\n"
            "Managing Director\t\tWhole-Time Director\n"
            "DIN: __________\t\tDIN: __________\n\n"
            "Date: ____________\n"
            "Place: ___________\n"
        )

        # Add the content to the document
        doc.add_paragraph(content)
        # Save the document
        # doc.save('Directors_Report.docx')

        filepath="board_report.docx"

        company_name_under= re.sub(r'\s+', '_', company_name)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

                    # AWS S3 configuration
        AWS_ACCESS_KEY = 'AKIAW2CLDMGVULENQAVM'
        AWS_SECRET_KEY = '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'
        BUCKET_NAME = 'legalnitiai'
        s3 = boto3.client('s3', aws_access_key_id=AWS_ACCESS_KEY, aws_secret_access_key=AWS_SECRET_KEY)

        file_name = f"board_report_{company_name_under}_{formatted_time}.docx"
        filepath = f"{file_name}"
                    
        heading = doc.add_heading("Board Report", level=1)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=attendance_sheet.docx'
                    # file_path = f"directors_table_{directors_count}_directors.docx"
                    
        s3.upload_fileobj(buffer, BUCKET_NAME, filepath, ExtraArgs={
                        # 'ContentType': 'application/pdf',
                        'ACL': 'public-read'
                    })


        print("File uploaded successfully to S3 bucket")

                    # Generate a pre-signed URL for the uploaded file
        full_url = f"https://{BUCKET_NAME}.s3.amazonaws.com/{filepath}"
        print("Full URL:", full_url)

                
        return Response({'success': True, 'file_path': full_url})     






        
                            
                


            





        

                        
