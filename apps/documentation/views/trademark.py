from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.views import View
import requests
import json
import urllib.parse
from django.http import HttpResponseBadRequest
import io
import urllib.request
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
from PIL import Image
from easyocr import Reader
from ultralytics.models.yolo import YOLO
from django.shortcuts import render
from django.http import HttpResponse
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document
from rest_framework.views import APIView
from rest_framework.response import Response
from django.views.generic import TemplateView
from docx.shared import Inches

from docx.shared import Pt
import openpyxl
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import requests
from rest_framework import status
import boto3
from botocore.exceptions import NoCredentialsError
from io import BytesIO
from datetime import datetime
import pytz
import os
import re


class trademark(APIView):

    def post(self,request,*args,**kwargs):

        ceo=request.data.get('name_ceo')
        mark=request.data.get('trademark') #only for the name 
        address=request.data.get('address')
        residence=request.data.get('residence')
        place=''
        date=request.data.get('date')
        company_name=request.data.get('company_name')



        current_datetime = datetime.now()

        # Format the current date and time
        formatted_time = current_datetime.strftime("%Y-%m-%d_%H-%M-%S")

        
        # Create a new document
        doc = Document()

        # Add the centered heading "TRADEMARK USER AFFIDAVIT"
        heading = doc.add_heading("TRADEMARK USER AFFIDAVIT", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        
        # Add the content paragraphs
        content_paragraphs = [
            f"In support of the request made to apply for Trademark for the mark “{mark}” under Trademark Class “30″.",
            f"I, {ceo} CEO of {mark} having principal place of business at {address}, do hereby solemnly affirm and declare as follows:",
            f"That I am Indian by nationality and residing at {residence}.",
            f"I state that I am familiar and well conversant with the facts and circumstances of the present matter and competent and authorized to swear this affidavit and make the necessary statement in respect thereof.",
            f"A trademark application is hereby made for registration of the accompanying trademark in class (30) and the said mark has been continuously used since 12th April 2023 in respect of the said Custom Manufacturing services.",
            f"That our company is engaged in the business of Custom Manufacturing services included in the Treatment of Materials which are provided under the said Trademark {mark}.",
            f"That the said Trademark {mark} was conceived and adopted by our firm and has been continuously used since 12th April 2023 in connection with the said business, and by reason of such use, the said business bearing the Trademark {mark} have come to be understood as being business provided by our company.",
            "That the trademark THE ELEVEN CAKES AND MORE has acquired a very valuable goodwill and reputation due to the sales and publicity and good quality of the Custom Manufacturing services.",
            f"That the trademark {mark} above respect to Custom Manufacturing services is solely and exclusively associated with the applicant and none else and has become distinctive on account of long continuous and extensive use and is entitled to registration.",
            "",
            "I solemnly state that the contents of this affidavit are true to the best of my knowledge and belief and that it conceals nothing and that no part of it is false.",
            "",
            "Signature",
            f"Place: {place}",
            f"Date: {date} ",
        ]

        for paragraph_text in content_paragraphs:
            paragraph = doc.add_paragraph(paragraph_text)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Add the centered heading "VERIFICATION"
        verification_heading = doc.add_heading("VERIFICATION", level=1)
        verification_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER


        # Add the signature, place, and date
        verification_signature = doc.add_paragraph("Signature")
        verification_signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        verification_place_date = doc.add_paragraph(f"Place: {place}\nDate: {date}")
        verification_place_date.alignment = WD_ALIGN_PARAGRAPH.LEFT


        # Save the document
    
        filepath="Trademark_User_Affidavit.docx"

        company_name_under= re.sub(r'\s+', '_', company_name)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

                    # AWS S3 configuration
        AWS_ACCESS_KEY = 'AKIAW2CLDMGVULENQAVM'
        AWS_SECRET_KEY = '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'
        BUCKET_NAME = 'legalnitiai'
        s3 = boto3.client('s3', aws_access_key_id=AWS_ACCESS_KEY, aws_secret_access_key=AWS_SECRET_KEY)

        file_name = f"trademark_{company_name_under}_{formatted_time}.docx"
        filepath = f"{file_name}"
                    
        heading = doc.add_heading("trademark", level=1)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=trademark.docx'
                    # file_path = f"directors_table_{directors_count}_directors.docx"
                    
        s3.upload_fileobj(buffer, BUCKET_NAME, filepath, ExtraArgs={
                        # 'ContentType': 'application/pdf',
                        'ACL': 'public-read'
                    })


        print("File uploaded successfully to S3 bucket")

                    # Generate a pre-signed URL for the uploaded file
        full_url = f"https://{BUCKET_NAME}.s3.amazonaws.com/{filepath}"
        print("Full URL:", full_url)

                
        return Response({'success': True, 'file_path': full_url})
