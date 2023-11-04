from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.views import View
import requests
import json
import urllib.parse
from django.http import HttpResponseBadRequest
import io
import urllib.request
import tempfile
from PIL import Image
from easyocr import Reader
from ultralytics.models.yolo import YOLO
from django.shortcuts import render
from django.http import HttpResponse
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document
from rest_framework.views import APIView
from rest_framework.response import Response
from django.views.generic import TemplateView
from docx import Document
from docx.shared import Pt
import openpyxl
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import requests
from rest_framework import status
import boto3
from botocore.exceptions import NoCredentialsError
from io import BytesIO
from datetime import datetime
import pytz
import datetime
from docx2pdf import convert
from docx.shared import Inches
import os
import re


@csrf_exempt
def AttendanceSheet(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        type_of_meeting = data.get('category["value"]')
        date_of_meeting = data.get('Date')
        time_of_meeting = data.get('Time')
        company_name = data.get('companyname')
        chairman_name = data.get('Chairman')
        # chairman_din = data.get('chairman_din')
        place_of_meeting = data.get('address')
        directors_count = int(data.get('selectedDirectorsLength', 0))

        current_datetime = datetime.datetime.now()

        # Format the current date and time
        formatted_time = current_datetime.strftime("%Y-%m-%d_%H-%M-%S")

        try:
            # Create a new document
            doc = Document()

            heading = doc.add_heading("Attendance Sheet", level=1)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            print("Date of Meeting:", date_of_meeting)
            print("Chairman Name:", chairman_name)
            print("Directors Count:", directors_count)
            print("Place of Meeting:", place_of_meeting)

            # Create the table for director details
            directors_table = doc.add_table(rows=1, cols=4)
            directors_table.style = "Table Grid"

            # Set the column widths
            column_widths = [1, 3, 3, 3]
            for i, width in enumerate(column_widths):
                directors_table.cell(0, i).width = width

            # Add the column headers for director details
            directors_headers = directors_table.rows[0].cells
            directors_headers[0].text = "Sr. No."
            directors_headers[1].text = "Name of Directors"
            directors_headers[2].text = "Designation"
            directors_headers[3].text = "DIN"

            for i, director_data in enumerate(data.get('selectedDirectors', [])):
                sr_no = str(i + 1)
                name = director_data.get('NAME', '')
                designation = "Director"
                din_number = director_data.get('DIN', '')

                row = directors_table.add_row().cells
                row[0].text = sr_no
                row[1].text = name
                row[2].text = designation
                row[3].text = din_number

                # Add designation and DIN number in the same cell
                # cell_3_paragraph = row[2].paragraphs[0]
                # cell_3_paragraph.text = designation
                # cell_3_paragraph.add_run("\n")  # Add a newline between designation and DIN number
                # cell_3_paragraph.add_run(din_number)

            # Add a blank paragraph for spacing
            doc.add_paragraph()

            # Add the company's name and chairman's details
            paragraph = doc.add_paragraph()
            paragraph.add_run("For and on behalf of the Board of Directors,").bold = True
            paragraph.add_run("\n" + company_name)

            # Add a blank paragraph for spacing
            doc.add_paragraph()

            # Add the line for signature
            doc.add_paragraph("_____________________________")

            # Add the chairman's details
            doc.add_paragraph(chairman_name)
            doc.add_paragraph("Chairman of the meeting")
            # doc.add_paragraph(f"DIN: {chairman_din}")

            # Add a blank paragraph for spacing
            doc.add_paragraph()

            # Add the date and place of the meeting
            doc.add_paragraph(f"Time: {time_of_meeting}")
            doc.add_paragraph(f"Date: {date_of_meeting}")
            doc.add_paragraph(f"Place: {place_of_meeting}")

            # Save the document
            # doc.save("attendance_sheet.docx")
            filepath = "attendance_sheet.docx"

            # pdf_file_path = filepath.replace(".docx", ".pdf")
            # convert(filepath, pdf_file_path)
            # ... Other imports ...

            company_name_under = re.sub(r'\s+', '_', company_name)

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            # AWS S3 configuration
            AWS_ACCESS_KEY = 'AKIAW2CLDMGVULENQAVM'
            AWS_SECRET_KEY = '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'
            BUCKET_NAME = 'legalnitiai'
            s3 = boto3.client('s3', aws_access_key_id=AWS_ACCESS_KEY, aws_secret_access_key=AWS_SECRET_KEY)

            file_name = f"attendance_sheet_{company_name_under}_{formatted_time}.docx"
            filepath = f"{file_name}"

            heading = doc.add_heading("Attendance Sheet", level=1)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = 'attachment; filename=attendance_sheet.docx'
            # file_path = f"directors_table_{directors_count}_directors.docx"

            s3.upload_fileobj(buffer, BUCKET_NAME, filepath, ExtraArgs={
                # 'ContentType': 'application/pdf',
                'ACL': 'public-read'
            })

            print("File uploaded successfully to S3 bucket")

            # Generate a pre-signed URL for the uploaded file
            full_url = f"https://{BUCKET_NAME}.s3.amazonaws.com/{filepath}"
            print("Full URL:", full_url)

            return JsonResponse({'success': True, 'file_path': full_url})

        except Exception as e:
            print("Error:", e)
            return JsonResponse({'success': False, 'error_message': str(e)})
    else:
        return JsonResponse({"message": "Invalid request method"}, status=400)

# ... Other views or code ...