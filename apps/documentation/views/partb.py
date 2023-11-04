from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.views import View
import requests
import json
import urllib.parse
from django.http import HttpResponseBadRequest
import io
import urllib.request
import tempfile
from PIL import Image
from easyocr import Reader
from ultralytics.models.yolo import YOLO
from django.shortcuts import render
from django.http import HttpResponse
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document
from rest_framework.views import APIView
from rest_framework.response import Response
from django.views.generic import TemplateView
from docx.shared import Inches

from docx.shared import Pt
import openpyxl
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import requests
from rest_framework import status
import boto3
from botocore.exceptions import NoCredentialsError
from io import BytesIO
from datetime import datetime
import pytz
import os
import re

@csrf_exempt
def partb(request):
    # Extract data from the request
    if request.method == 'POST':
        data = json.loads(request.body)
        name = data.get('name')
        father_name = data.get('father_name')
        address = data.get('address')
        company_name = data.get('company_data')

        # Get the current date and time
        current_datetime = datetime.now()

        # Format the current date and time
        formatted_time = current_datetime.strftime("%Y-%m-%d_%H-%M-%S")

        # Create a new Word document
        doc = Document()

        # Add a centered heading to the document (level 1)
        heading = doc.add_heading("Part-B: Statement", level=1)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add a centered heading of level 2
        heading2 = doc.add_heading("Statement by a person who subscribed his name to the incorporation document", level=2)
        heading2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add content to the document
        content = (
            f"I, {name} , S/o {father_name}, R/o, {address}, "
            f"the Designated Partner of the {company_name} do state that:\n\n"
            "(i) I am a person named in the incorporation document as a Designated Partner/partner of the limited liability partnership;\n"
            "(ii) The Designated Partner have given their prior consent to act as a Designated Partner;\n"
            "(iii) All the requirements of the Limited Liability Partnership Act, 2008 and rules made thereunder in respect of designated partner identification number (DPIN), "
            "registration of the LLP, and matters precedent and incidental thereto have been complied with;\n"
            "(iv) I make this statement conscientiously believing the same to be true.\n\n"
            "Thank You\n"
        )

        content_paragraph = doc.add_paragraph(content)
        content_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Add "Yours sincerely,"
        sincerely_text = "Yours sincerely,"
        sincerely_paragraph = doc.add_paragraph(sincerely_text)
        sincerely_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Add an image with signature aligned to the left bottom
        signature_text = "Signature"
        signature_paragraph = doc.add_paragraph(signature_text)
        signature_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Add additional information
        additional_info = f"{name}\n"
        additional_info_paragraph = doc.add_paragraph(additional_info)
        additional_info_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Generate a file name based on company name and timestamp
        company_name_under = re.sub(r'\s+', '_', company_name)
        file_name = f"partb_{company_name_under}_{formatted_time}.docx"

        # Save the document to a buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # AWS S3 configuration
        AWS_ACCESS_KEY = 'AKIAW2CLDMGVULENQAVM'
        AWS_SECRET_KEY = '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'
        BUCKET_NAME = 'legalnitiai'
        s3 = boto3.client('s3', aws_access_key_id=AWS_ACCESS_KEY, aws_secret_access_key=AWS_SECRET_KEY)

        # Upload the document to S3
        s3.upload_fileobj(buffer, BUCKET_NAME, file_name, ExtraArgs={'ACL': 'public-read'})

        # Generate a pre-signed URL for the uploaded file
        full_url = f"https://{BUCKET_NAME}.s3.amazonaws.com/{file_name}"

        # Return the response with the URL
        return JsonResponse({'success': True, 'file_path': full_url})
    else:
        return JsonResponse({"message": "Invalid request method"}, status=400)
