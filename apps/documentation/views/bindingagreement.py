from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.views import View
import requests
import json
import urllib.parse
from django.http import HttpResponseBadRequest
import io
import urllib.request
import tempfile
from PIL import Image
from easyocr import Reader
from ultralytics.models.yolo import YOLO
from django.shortcuts import render
from django.http import HttpResponse
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document
from rest_framework.views import APIView
from rest_framework.response import Response
from django.views.generic import TemplateView
from docx import Document
from docx.shared import Pt
import openpyxl
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import requests
from rest_framework import status
import boto3
from botocore.exceptions import NoCredentialsError
from io import BytesIO
from datetime import datetime
import pytz
import datetime
from docx2pdf import convert
from docx.shared import Inches
import os
import re


class bindingagreement(APIView):

    def post(self,request,*args,**kwargs):

        date=request.data.get('date')
        residence=request.data.get('residence')
        proprietor=request.data.get('proprietor',[])
        partyc=request.data.get('partyc',[])
        partyd=request.data.get('partyd',[])
        company_name=request.data.get('companyname','')

        #propreitor
        proprietor_name=proprietor.get('name','')
        proprietor_pan=proprietor.get('pan','')
        proprietor_residence=proprietor.get('residence','')

        #partyc
        partyc_name=partyc.get('name','')
        partyc_father_name=partyc.get('father_name','')
        partyc_pan=partyc.get('pan','')
        partyc_residence=partyc.get('residence','')

        #partyd
        partyd_name=partyd.get('name','')
        partyd_father_name=partyd.get('father_name','')
        partyd_pan=partyd.get('pan','')
        partyd_residence=partyd.get('residence','')


        current_datetime = datetime.datetime.now()

        # Format the current date and time
        formatted_time = current_datetime.strftime("%Y-%m-%d_%H-%M-%S")



        doc = Document()

        # Add the specified content
        content = [
            f"This Multi-Party Agreement (\"Agreement\") is officially established and enacted on the {date}, hereinafter referred to as the \"Effective Date.\"",
            "The Agreement, herein referred to as the \"Binding Agreement,\" stands as a tangible embodiment and binding representation of the intentions, commitments, and shared aspirations of the entities listed below. This Agreement, is a legally binding instrument, solidifies the terms, conditions, and mutual obligations that define the relationship between the Parties:",
            f"Soul for Sole(SFS), Sole Proprietorship (yet to be incorporated under the laws of Companies Act, 2013,) and having its registered office at {residence}, India hereinafter referred to as Party A;",
            "AND",
            f"{proprietor_name}, sole proprietor, D/o Kishor Gawali having PAN number {proprietor_pan} and residing at {proprietor_residence} under whose name Soul for Sole is being registered for any kind of registration required by the Proprietorship like GST registration, Trademark registration etc, hereinafter referred to as Party B;",
            "AND",
            f"{partyc_name}, individual, S/o {partyc_father_name} having PAN number {partyc_pan} and residing at {partyc_residence} hereinafter referred to as Party C;",
            "AND",
            f"{partyd_name}, individual, S/o {partyd_father_name} having PAN number {partyd_pan} and residing at {partyd_residence} hereinafter referred to as Party D;",
            f"The Soul For Sole(SFS) and {proprietor_name}, {partyc_name}, {partyd_name} are hereinafter individually referred to as a Party and collectively as the Parties.",
            "Now, Therefore, in consideration of the promises and the mutual covenants set forth herein, the Parties hereto, intending to be legally bound, hereby agree as follows:",
            "WHEREAS, the Parties acknowledge and recognize that the proprietorship \"Soul for Sole\" is currently registered under the name of Samiksha Kishor Gawali , but the liability of the proprietorship belongs to all Parties equally;",
            "WHEREAS, the Parties share a common vision and dedication to the growth and prosperity of \"Soul for Sole,\" and their collective efforts are intended to contribute to the proprietorship's success;",
            "WHEREAS, the Parties intend to formalize their joint ownership and mutual commitment to the proprietorship's present and future endeavors;",
            "WHEREAS, the Parties are committed to fostering a collaborative and transparent partnership that reflects the principles of shared responsibility, shared decision-making, and shared benefits;",
            "WHEREAS, the Parties recognize the need to outline the rights, responsibilities, and expectations of each Party in relation to their roles within \"Soul for Sole\";",
            "NOW, THEREFORE, in consideration of the mutual promises and covenants contained herein, the Parties hereby agree as follows:",
            "1. Purpose and scope:",
            "1.1 Soul for Sole(SFS) (\"Sole Proprietorship\") is in the process of being registered under the legal name of Party A .",
            "1.2 The Parties acknowledge and agree that despite the Proprietorship being registered under Party B’s name, the Proprietorship's operations and business matters will be governed by the terms of this Agreement.",
            "1.3 The Parties understand that the registration is in the name of Party B solely for legal and administrative purposes and does not alter the shared ownership, responsibilities, or liabilities as set forth in this Agreement.",
            "1.4 The purpose of this Agreement is to establish a clear framework for the collaboration among the Parties, outlining their collective ownership, joint decision-making, and mutual commitment to the success and growth of the Proprietorship.",
            "1.5 This Agreement serves to formalize the relationship among the Parties and ensure that their contributions, both individually and collectively, are recognized and valued in the proprietor's operations and future endeavors.",
            "1.6 The Parties intend for this Agreement to act as a foundational document that reflects their shared vision, values, and aspirations for Party A, transcending individual interests and fostering a unified approach to achieving business objectives.",
            "1.7 Through this Agreement, the Parties seek to create a cohesive and enduring partnership that not only addresses present circumstances but also safeguards the Proprietorship's evolution, adaptability, and resilience to changes in the market, technology, and business landscape.",
            "2. Ownership and Decision-Making:",
            "2.1 The Parties, commit to a proportional division of ownership in the ratio of 1:1:1 within the Proprietorship.",
            "2.2 Recognizing the significance of unity and cooperation, all major decisions intricately shaping the trajectory of the Proprietorship's operations, strategic investments, potential partnerships, and other pivotal matters shall be undertaken collectively. This collective decision-making approach exemplifies the Parties' unwavering commitment to fostering an environment of inclusivity, where each voice is equally respected and each perspective is thoughtfully considered.",
            "2.3 The inherent strength of this collective decision-making process lies in its ability to harness the diverse expertise, insights, and visions brought forth by the parties. By joining forces in this manner, the Parties are poised to leverage their combined knowledge, skills, and strategic acumen for the betterment of the Proprietorship's growth and prosperity.",
            "2.4 This collaborative decision-making mechanism ensures that no single Party dominates the discourse, reinforcing a balanced power dynamic that promotes transparency, accountability, and the ethical pursuit of shared objectives. The Parties' commitment to consult, deliberate, and agree on vital matters underscores their determination to nurture a culture of trust and consensus within the Proprietorship.",
            "2.5 Through this collective approach, the parties embark on a shared journey, where the Proprietorship's direction and evolution are shaped by a symphony of perspectives, guided by a deep-rooted ethos of joint responsibility and shared outcomes.",
            "2.6 The Parties firmly believe that this collaborative decision-making framework not only fortifies their collective resolve but also serves as a cornerstone for the sustainable success and enduring legacy of the Proprietorship, creating a legacy that resonates far beyond individual milestones.",
            "3. Proprietorship’s Responsibility:",
            "3.1 Devise comprehensive marketing strategies and campaigns that effectively promote the Proprietorship's products to target audiences.",
            "3.2 Execute creative and result-oriented marketing initiatives, encompassing digital and traditional channels, to enhance brand visibility and customer engagement.",
            "3.3 Conduct market research and analysis to identify consumer trends, preferences, and competitive landscape, informing the Proprietorship's marketing decisions.",
            "3.4 Collaborate with external partners, agencies, and stakeholders to optimize marketing efforts and ensure consistent brand messaging.",
            "4. Party A, Party B and Party C- Joint Responsibility:",
            "4.1 Collaboratively develop and execute comprehensive marketing strategies, campaigns, and initiatives to enhance brand visibility and customer engagement.",
            "4.2 Coordinate operational aspects, including inventory management, procurement, logistics, and retail, to ensure efficient product delivery and customer satisfaction.",
            "4.3 Foster a cohesive and harmonious work environment by leveraging individual strengths, skills, and insights to drive the Proprietorship's success.",
            "4.4 Implement financial management practices, including budgeting, financial planning, and compliance, to ensure prudent financial decisions and sustainable growth",
            "5. Liabilities, Profits, and Losses:",
            "5.1 Despite the registration of the Proprietorship under Party B’s name, all Parties shall share the liabilities, profits, losses, and responsibilities equally, in proportion to their ownership shares.",
            "5.1.1 Embrace a collective ownership approach, recognizing that while registered under Party B’s name, the Proprietorship's essence lies in shared responsibility and equal distribution.",
            "5.1.2 Affirm the commitment to distribute all outcomes, whether financial gains, losses, or responsibilities, in a balanced manner reflective of each Parties ownership share.",
            "5.1.3 Understand that this equal sharing principle extends across all spheres of the Proprietorship, ensuring that all Parties contribute and benefit in an integrated manner.",
            "5.1.4  Acknowledge that this ethos empowers all Parties to actively participate in decision-making, leveraging diverse perspectives for collective growth.",
            "5.1.5 Emphasize that equal distribution signifies a readiness to stand united in both prosperous and challenging times, reinforcing a shared path forward.",
            "5.1.6 View this approach as a foundation for an aligned responsibility framework, where ownership extends beyond mere registration to inclusive participation.",
            "5.1.7 Recognize that equal sharing fosters transparent accountability, promoting honest communication and trust among all Parties involved.",
            "5.1.8 Understand that this principle fortifies the Proprietorship's strategic unity, enabling cohesive actions that leverage collective strengths for sustainable success.",
            "5.1.9 Equal sharing paradigm, elevate the collective impact of all Parties, propelling the Proprietorship toward enduring accomplishments.",
            "5.2 The Parties acknowledge and agree to contribute equally to any financial requirements of the Proprietorship.",
            "5.2.1 Reiterate the shared pledge of all Parties to provide equal financial contributions, solidifying a collaborative commitment to the Proprietorship's fiscal health.",
            "5.2.2 Recognize that financial requirements encompass diverse aspects, encompassing operational expenditures, expansions, innovations, and unforeseen needs.",
            "5.2.3 Emphasize that equal financial contribution guarantees balanced participation in the Proprietorship's financial journey, upholding a fair and inclusive partnership.",
            "5.2.4 Understand that this unified financial support empowers the Proprietorship's growth, enabling well-informed decisions driven by collective resources.",
            "5.2.5 View equal financial contribution as strategic resource allocation, amplifying the Proprietorship's capacity to make impactful investments aligned with shared goals.",
            "5.2.6 Foster a sense of reliability by upholding equal financial contribution, underscoring a dependable partnership characterized by consistent support.",
            "5.2.7 Align the Parties' financial commitments with the Proprietorship's vision, illustrating a synchronized pursuit of financial objectives for mutual prosperity.",
            "5.3 Future decisions and ownership claim:",
            "5.3.1 The Parties recognize and acknowledge that the business landscape is a dynamic entity, subject to continual evolution, and liable to present both new avenues of opportunities and unforeseen challenges.",
            "5.3.2 In the unlikely event that any Parties assert sole ownership claims in the future, this Agreement shall stand as an unwavering bastion, firmly establishing the bedrock of collective ownership, shared commitments, and explicitly defined responsibilities.",
            "5.3.3 The comprehensive nature of this Agreement ensures that any potential claims of sole ownership, no matter how improbable, shall not hold the power to modify or reshape the pre-determined distribution of ownership interests and the corresponding obligations as meticulously detailed herein.",
            "5.3.4 The Agreement's authority transcends the realm of hypothetical claims, upholding the sanctity of the Parties' collective understanding and agreement as the ultimate governing authority over matters of ownership and responsibilities.",
            "5.3.5 The Parties recognize that the strength of their collective commitment and the sanctity of this Agreement shall prevail over any individual assertions, ensuring the continuity of shared ownership, obligations, and responsibilities.",
            "5.3.6 Any prospective sole ownership claims shall be treated with the utmost regard for the collective interests of all Parties, with the understanding that the core principles established within this Agreement remain unalterable.",
            "5.3.7 The Parties recognize that the enduring nature of this Agreement transcends individual interpretations and upholds the fundamental tenets of shared ownership, thereby ensuring the holistic well-being and sustainability of the Proprietorship.",
            "5.3.8 The Agreement's resilience against any potential sole ownership claims demonstrates the Parties' shared commitment to a harmonious, collaborative, and unified journey, irrespective of potential deviations.",
            "5.3.9 The overarching goal of this Agreement is to create a framework where the collective spirit of ownership thrives, underlining the Parties' shared dedication to nurturing the Proprietorship's success through equal involvement and mutual trust.",
            "5.3.10 In essence, this Agreement serves as an unshakable pillar that fortifies the Parties' united ownership, resonating through potential challenges and triumphs, and solidifying the unbreakable bond that propels the Proprietorship's shared journey forward.",
            "6. Confidentiality Obligations:",
            "6.1 All Parties shall diligently uphold the imperative of maintaining unwavering confidentiality concerning proprietary and sensitive information related to the Proprietorship's operations and business strategies.",
            "6.2 The commitment to confidentiality remains in effect even after the termination of this Agreement.",
            "6.3 The Parties pledge to strictly guard the Proprietorship's proprietary information, ensuring its non-disclosure to unauthorized entities.",
            "6.4 This commitment signifies not only a legal obligation but also a fundamental aspect of trust and ethical business conduct.",
            "6.5 By extending the obligation beyond the Agreement's conclusion, the Parties reaffirm their dedication to safeguarding the Proprietorship's interests and preserving its integrity.",
            "7. Duration and Proprietorship Incorporation:",
            "7.1 This Agreement shall maintain its validity until the Proprietorship undergoes the process of incorporation as a private limited entity.",
            "7.2 Upon the successful completion of the incorporation process, the provisions outlined in this Agreement shall seamlessly transition to become an integral part of the Proprietorship's new corporate framework.",
            "7.3 The inherent principles of shared ownership, collective responsibilities, and agreed-upon commitments shall persist within the Proprietorship's structure, ensuring the continuity of the Parties' unified journey.",
            "7.4 The Proprietorship's transition to a private limited entity shall stand as a testament to the enduring nature of this Agreement, which persists even through transformative milestones, upholding the Parties' shared vision.",
            "7.5 The harmonious integration of this Agreement into the Proprietorship's status exemplifies the Parties' unwavering dedication to a collaborative and united approach, further reinforcing the Proprietorship's foundation.",
            "8. Representations and Warranties:",
            "8.1 Each Party hereby represents and warrants that they have full legal capacity and authority to enter into this Agreement.",
            "8.2 Each Party acknowledges that they have not relied on any verbal or written representations, other than those expressly stated in this Agreement, in making their decisions to be bound by its terms.",
            "8.3 Each Party warrants the accuracy and completeness of any information, documentation, or materials they provide to the Proprietorship during this Agreement.",
            "8.4 Each Party represents that they are not under any legal impediment, contractual obligation, or conflict of interest that would hinder their faithful performance of the obligations under this Agreement.",
            "8.5 The Parties collectively affirm that they have conducted sufficient due diligence and obtained independent advice, as required, to comprehend the implications and ramifications of this Agreement.",
            "8.6 Each Party warrants that they will promptly notify the others of any material change in their legal status or capacity that may impact their ability to fulfill the obligations under this Agreement.",
            "8.7 The Parties warrant that they have not engaged in any fraudulent activity, misrepresentation, or concealment of material facts related to the negotiation or execution of this Agreement.",
            "9. Dispute Resolution:",
            "9.1 In the event of any dispute arising out of or in connection with this Agreement, the Parties shall first attempt to resolve the matter amicably through direct negotiations.",
            "9.2 If the dispute cannot be resolved, the Parties agree to submit the dispute to mediation before pursuing any legal action.",
            "9.3 The mediation process shall be conducted by a neutral and accredited mediator chosen by mutual agreement of the Parties.",
            "9.4 The Parties agree to participate in the mediation process in good faith, making genuine efforts to find a satisfactory solution.",
            "9.5 If mediation fails to produce a resolution, either Party may initiate legal proceedings, seeking remedies available under the relevant laws and jurisdictions.",
            "9.6 The Parties acknowledge that the pursuit of legal action is a secondary recourse, and they shall strive to minimize adversarial proceedings through mediation.",
            "9.7 Any resolutions reached through mediation shall be binding upon the Parties, and they shall undertake to execute any necessary actions to implement such resolutions.",
            "10. Compliance with Laws:",
            "10.1 The Parties commit to conducting all Proprietorship activities in strict accordance with prevailing laws and regulations.",
            "10.2 The Parties collectively ensure that the Proprietorship consistently adheres to legal requirements to maintain ethical and lawful operations.",
            "10.3 Regular reviews will be conducted to ensure the Proprietorship's activities remain compliant with evolving legal standards.",
            "10.4 Should any legal uncertainties arise; the Parties will seek professional legal counsel for accurate guidance.",
            "10.5 Comprehensive records of legal matters, transactions, and compliance efforts will be maintained to ensure transparency and accountability.",
            "11. Governing Law and Jurisdiction:",
            "11.1 This Agreement shall be governed by and construed in accordance with the laws of applicable jurisdiction.",
            "11.2 Any disputes arising from or related to this Agreement shall be subject to the exclusive jurisdiction of the courts in applicable jurisdiction.",
            "15. Entire Agreement and Liability:",
            "15.1 The Parties acknowledge that this Agreement constitutes the entire understanding between them regarding the subject matter and supersedes all previous communications, representations, or agreements, whether written or oral.",
            "15.2 This Agreement, along with its annexes or attachments, shall constitute a legally binding document that governs the Parties' relationship and responsibilities.",
            "15.3 Any modifications, amendments, or changes to this Agreement shall be valid only if made in writing and duly executed by all Parties.",
            "15.4 The Parties acknowledge that each Party shall be jointly and severally liable for the obligations, representations, and commitments set forth in this Agreement.",
            "15.5 Each Party agrees to indemnify and hold harmless the other Parties from any claims, losses, damages, liabilities, or expenses arising out of any breach or failure to fulfil the obligations outlined in this Agreement.",
            "15.6 Except for instances of wilful misconduct or gross negligence, none of the Parties shall be liable to each other for any indirect, special, consequential, or punitive damages arising from the performance or non-performance of this Agreement.7.6 Should any provision of this Agreement be found unenforceable or invalid, it shall not affect the enforceability or validity of the remaining provisions, which shall continue to bind the Parties.",
            "15.7 The failure of any Party to enforce any provision of this Agreement shall not constitute a waiver of that provision or the right to subsequently enforce it.",
            "15.8 This Agreement has been duly executed by the authorized representatives of the Parties as of the Effective Date."
            
            "IN WITNESS WHEREOF, the Parties hereto have executed this Multi-Party Agreement as of the Effective Date.",
            "Party A: ________________________",
            "Name: ________________________",
            "Party B: ________________________",
            "Name: ________________________",
            "Party C: ________________________",
            "Name: ________________________",
            "Party D: ________________________",
            "Name: ________________________"
        ]

        # Add the content paragraphs to the document
        for paragraph in content:
            doc.add_paragraph(paragraph)

        content = [
        
            "Authorized Signatory                  Authorized Signatory",
            f"{proprietor_name}                                         "
        ]

        # Add the content paragraphs to the document
        for paragraph in content:
            doc.add_paragraph(paragraph)


        # Save the document
        

        filepath="binding_agreement.docx"

        company_name_under= re.sub(r'\s+', '_', company_name)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

                # AWS S3 configuration
        AWS_ACCESS_KEY = 'AKIAW2CLDMGVULENQAVM'
        AWS_SECRET_KEY = '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'
        BUCKET_NAME = 'legalnitiai'
        s3 = boto3.client('s3', aws_access_key_id=AWS_ACCESS_KEY, aws_secret_access_key=AWS_SECRET_KEY)

        file_name = f"binding_agreement_{company_name_under}_{formatted_time}.docx"
        filepath = f"{file_name}"
                    
        heading = doc.add_heading("binding agreement", level=1)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=binding_agreement'
                    # file_path = f"directors_table_{directors_count}_directors.docx"
                    
        s3.upload_fileobj(buffer, BUCKET_NAME, filepath, ExtraArgs={
                        # 'ContentType': 'application/pdf',
                        'ACL': 'public-read'
                    })


        print("File uploaded successfully to S3 bucket")

                    # Generate a pre-signed URL for the uploaded file
        full_url = f"https://{BUCKET_NAME}.s3.amazonaws.com/{filepath}"
        print("Full URL:", full_url)

                
        return Response({'success': True, 'file_path': full_url})