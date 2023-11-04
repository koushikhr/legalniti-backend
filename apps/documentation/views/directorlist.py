from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.views import View
import requests
import json
import urllib.parse
from django.http import HttpResponseBadRequest
import io
import urllib.request
import tempfile
from PIL import Image
from easyocr import Reader
from ultralytics.models.yolo import YOLO
from django.shortcuts import render
from django.http import HttpResponse
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document
from rest_framework.views import APIView
from rest_framework.response import Response
from django.views.generic import TemplateView
from docx import Document
from docx.shared import Pt
import openpyxl
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import requests
from rest_framework import status
import boto3
from botocore.exceptions import NoCredentialsError
from io import BytesIO
from datetime import datetime
import pytz
import datetime
from docx2pdf import convert
from docx.shared import Inches
import os
import re

class DirectorList(APIView):
    def post(self, request, *args, **kwargs):

        num_directors = int(request.data.get('num_directors', 0))

        if num_directors < 1:
            return Response({'error': 'Invalid input. Please enter a positive integer.'}, status=400)

        doc = Document()
        heading_paragraph = doc.add_paragraph()
        heading_run = heading_paragraph.add_run("LIST OF DIRECTORS")
        heading_run.bold = True
        # heading_run.font.size = Pt(14)
        heading_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'

        headers = [request.data.get('din_header', 'DIN/PAN/DPIN'),
                   request.data.get('name_header', 'Full Name'),
                   request.data.get('designation_header', 'Designation'),
                   request.data.get('address_header', 'Address')]
        
        
       
        header_cells = table.rows[0].cells
        for i in range(len(headers)):
            header_cells[i].text = headers[i]

        for i in range(num_directors):
            row = table.add_row().cells
            row[0].text = request.data.get(f'din_{i}', '')
            row[1].text = request.data.get(f'full_name_{i}', '')
            row[2].text = request.data.get(f'designation_{i}', '')
            row[3].text = request.data.get(f'address_{i}', '')

        file_path = f"directors_table_{num_directors}_directors.docx"
        doc.save(file_path)

        return Response({'success': True, 'file_path': file_path})




class Minutemeet(APIView):
    def post(self, request, *args, **kwargs):
        data = request.data

        # Extract data from the request
        
        company_name = data.get('companyname')
        date_of_meeting = data.get('Date')
        location_of_meeting = data.get('address')
        directors_count = int(data.get('selectedDirectorsLength',0))


        current_datetime = datetime.datetime.now()

        # Format the current date and time
        formatted_time = current_datetime.strftime("%Y-%m-%d_%H-%M-%S")



        # Create a new document
        doc = Document()

        # doc.add_heading(f"Board Meeting - No. {board_meeting_number}", level=1)

        minutes_text = f"MINUTES OF THE MEETING OF THE BOARD OF DIRECTORS OF {company_name}"
        minutes_text += f" HELD ON {date_of_meeting} AT THE REGISTERED OFFICE OF THE COMPANY SITUATED AT {location_of_meeting}"

        # Add the first paragraph
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(minutes_text)
        run.bold = True

        # Add a heading after "DIRECTORS PRESENT" paragraph
        doc.add_heading("DIRECTORS PRESENT", level=2)

        # Create the table with 4 columns
        # Create the table for director details
        directors_table = doc.add_table(rows=1, cols=4)
        directors_table.style = "Table Grid"

            # Set the column widths
        column_widths = [1, 3, 3, 3]
        for i, width in enumerate(column_widths):
                directors_table.cell(0, i).width = width

            # Add the column headers for director details
        directors_headers = directors_table.rows[0].cells
        directors_headers[0].text = "Sr. No."
        directors_headers[1].text = "Name of Directors"
        directors_headers[2].text = "Designation"
        directors_headers[3].text = "DIN"

        # Prompt the user for director details and add them to the table
        for i, director_data in enumerate(data.get('selectedDirectors', [])):
                sr_no = str(i + 1)
                name = director_data.get('NAME', '')
                designation = "Director"
                din_number = director_data.get('DIN', '')

                row = directors_table.add_row().cells
                row[0].text = sr_no
                row[1].text = name
                row[2].text = designation
                row[3].text = din_number

        # Add a heading for "Period of Meeting"
        doc.add_heading("Period of Meeting", level=2).bold = True

        # Prompt the user for time of commencement and conclusion
        time_of_commencement = data.get("Start Time")
        time_of_conclusion = data.get("End Time")

        # Add the values for time of commencement and conclusion
        paragraph = doc.add_paragraph()
        paragraph.add_run("Time of Commencement: ").bold = True
        paragraph.add_run(time_of_commencement)

        paragraph = doc.add_paragraph()
        paragraph.add_run("Time of Conclusion: ").bold = True
        paragraph.add_run(time_of_conclusion)

        # Add a heading for "Agenda"
        doc.add_heading("Agenda", level=2).bold = True

        # Initialize the agenda counter
        agenda_counter = 1


        agenda_items = data.get('agenda_items', []) 

        # Iterate through multiple agenda items
        for agenda_item_data in agenda_items:
            agenda_item = agenda_item_data.get('agenda_item')
            proposal = agenda_item_data.get('proposal')

            # Add the agenda item and the proposal to the document
            paragraph = doc.add_paragraph()
            paragraph.add_run(f"{agenda_counter}) {agenda_item}").bold = True
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            paragraph = doc.add_paragraph()
            paragraph.add_run(proposal)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            # Increment the agenda counter
            agenda_counter += 1
                # Add page numbers to the document
        sections = doc.sections
        for section in sections:
            footer = section.footer

            # Get or create the footer paragraph
            if footer is None:
                footer = section.footer
                footer_paragraph = footer.paragraphs[0]
            else:
                footer_paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

            # Clear existing content in the footer
            for elem in footer_paragraph._p:
                footer_paragraph._p.remove(elem)

            
            

            

        # Add Chairman of the meeting information and signature line
        chairman_name = data.get("Chairman")
        # chairman_din = data.get("Enter DIN of Chairman: ")

        doc.add_heading("Chairman of the Meeting", level=2).bold = True
        doc.add_paragraph(f"{chairman_name}")
        # doc.add_paragraph(f"DIN: {chairman_din}")
        doc.add_paragraph("")

        # Add signature line for the Chairman
        doc.add_paragraph("Signature: __________")

        # doc.save("minute_meet.docx")
        filepath="director_list.docx"


        company_name_under= re.sub(r'\s+', '_', company_name)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

            # AWS S3 configuration
        AWS_ACCESS_KEY = 'AKIAW2CLDMGVULENQAVM'
        AWS_SECRET_KEY = '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'
        BUCKET_NAME = 'legalnitiai'
        s3 = boto3.client('s3', aws_access_key_id=AWS_ACCESS_KEY, aws_secret_access_key=AWS_SECRET_KEY)

        file_name = f"director_list_{company_name_under}_{formatted_time}.docx"
        filepath = f"{file_name}"
            
        heading = doc.add_heading("Attendance Sheet", level=1)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=attendance_sh.docx'
            # file_path = f"directors_table_{directors_count}_directors.docx"
            
        s3.upload_fileobj(buffer, BUCKET_NAME, filepath, ExtraArgs={
                # 'ContentType': 'application/pdf',
                'ACL': 'public-read'
            })


        print("File uploaded successfully to S3 bucket")

            # Generate a pre-signed URL for the uploaded file
        full_url = f"https://{BUCKET_NAME}.s3.amazonaws.com/{filepath}"
        print("Full URL:", full_url)

           
        return Response({'success': True, 'file_path': full_url})

        

# ... Other views or code ...


        





