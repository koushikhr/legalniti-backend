from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.views import View
import requests
import json
import urllib.parse
from django.http import HttpResponseBadRequest
import io
import urllib.request
import tempfile
from PIL import Image
from easyocr import Reader
from ultralytics.models.yolo import YOLO
from django.shortcuts import render
from django.http import HttpResponse
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document
from rest_framework.views import APIView
from rest_framework.response import Response
from django.views.generic import TemplateView
from docx import Document
from docx.shared import Pt
import openpyxl
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import requests
from rest_framework import status
import boto3
from botocore.exceptions import NoCredentialsError
from io import BytesIO
from datetime import datetime
import pytz
import datetime
from docx2pdf import convert
from docx.shared import Inches
import os
import re



class GenerateNoticeView(APIView):
    
    def post(self, request, *args, **kwargs):
            
            type_of_meeting = request.data.get('category.value')
            # serial_number = request.data.get('serial_number')
            date_of_meeting = request.data.get('Date')
            time_of_meeting = request.data.get('Time')
            company_name = request.data.get('companyname')
            exact_address = request.data.get('address')
            # place_of_meet = request.data.get('place_of_meet')
            chairman_name = request.data.get('Chairman')
            director_data = request.data.get('selectedDirectors', [])
            DIN_number= director_data.get('DIN', '')
            director_name = director_data.get('NAME', '')
            
            
            current_datetime = datetime.datetime.now()

        # Format the current date and time
            formatted_time = current_datetime.strftime("%Y-%m-%d_%H-%M-%S")
            
            doc = Document()





            # Create the letterhead section
            #letterhead = doc.sections[0].header
            #letterhead.is_linked_to_previous = False

            # Add the heading to the letterhead
            #heading = letterhead.paragraphs[0].add_run("Letterhead")
            #heading.bold = False
            #heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Print the board meeting number
            # doc.add_paragraph(f"Board Meeting Number: {serial_number}")

            # Add the notice of the meeting
            notice_paragraph = doc.add_paragraph()
            notice_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # Add the bold "Notice of Meeting" text
            notice_run = notice_paragraph.add_run("Notice of Board Meeting")
            notice_run.bold = True
            notice_run.font.size = Pt(14)

            doc.add_paragraph("To,\n"
                            "The Board of Directors,\n"
                            f"{company_name},\n"
                            f"{exact_address}")

            # Add the salutation
            doc.add_paragraph("Respected Sir/Madam,")

            # Add the notice of meeting
            paragraph = doc.add_paragraph("Notice is hereby given that the meeting of the Board of Directors of the Company "
                                        + company_name + " "
                                        + "(the 'Board') scheduled to be held on " + date_of_meeting + " at " + time_of_meeting + " "
                                        + "at the registered office of the company situated " + exact_address + ", India "
                                        + "to transact the below-mentioned business:")


            agendacounter = 1


            # list of agenda
            for i,agenda_data in enumerate(request.data.get('agenda_items', [])):
                agenda_item = agenda_data.get('agenda','')
                

                # Add the agenda item and the proposal to the document
                paragraph = doc.add_paragraph()
                paragraph.add_run(f"{agendacounter}) {agenda_item}").bold = True
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            
                # Increment the agenda counter
                agendacounter += 1

            
            
            
            
            #Add the additional paragraph after the agenda list
            doc.add_paragraph("Kindly make it convenient to attend the meeting. Please submit leave of absence in case you are not in a position to attend the meeting.")

            # Add the additional paragraph after the agenda list
            additional_paragraph = doc.add_paragraph("By and on behalf of \n")
            additional_paragraph.add_run(company_name).bold = True
            additional_paragraph.add_run("\n\n_____________")
            additional_paragraph.add_run("\n\n")
            additional_paragraph.add_run(chairman_name)
            additional_paragraph.add_run("\nDirector")
            additional_paragraph.add_run("\nDIN: " + DIN_number)
            additional_paragraph.add_run("\n\nDate: " + date_of_meeting)
            additional_paragraph.add_run("\nPlace: " + exact_address)

            # Copy of the notice to:
            doc.add_paragraph("\nCopy of the notice to:\n")


            #enter the chairman of the meet
            # chairman_present = input("enter the chairman present ")
            # din_of_chairman = input("enter the din of the chairman ")
            # Add the NOTES TO AGENDA section
            notes_to_agenda = f"NOTES TO AGENDA: {type_of_meeting} DATED {date_of_meeting}:"
            notes_to_agenda_paragraph = doc.add_paragraph("\n" + notes_to_agenda)
            notes_to_agenda_paragraph.bold = True

            # Iterate through each agenda item and proposal
            # for agenda_text, proposal in agenda_list:
            #     item_paragraph = doc.add_paragraph("ITEM NO. " + str(agenda_counter) + ": " + agenda_text)
            #     item_paragraph.add_run("\n\n" + proposal)
                
            #     # Increment the agenda counter
            #     agenda_counter += 1

            
            doc.add_heading("Agenda", level=2).bold = True

            # Initialize the agenda counter
            

            agenda_counter=1
            

            # Iterate through multiple agenda items
            for i,agenda_data in enumerate(request.data.get('agenda_items', [])):
                agenda_item = agenda_data.get('agenda','')
                proposal = agenda_data.get('proposal','')

                # Add the agenda item and the proposal to the document
                paragraph = doc.add_paragraph()
                paragraph.add_run(f"{agenda_counter}) {agenda_item}").bold = True
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                paragraph = doc.add_paragraph()
                paragraph.add_run(proposal)
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                # Increment the agenda counter
                agenda_counter += 1




            # Add the "By and on behalf of" section
            by_on_behalf_paragraph = doc.add_paragraph("By and on behalf of\n")
            by_on_behalf_paragraph.add_run(company_name).bold = True
            by_on_behalf_paragraph.add_run("\n\n_____________")
            by_on_behalf_paragraph.add_run("\n\n")
            by_on_behalf_paragraph.add_run("Chairman of the meeting")
            by_on_behalf_paragraph.add_run("\n" + director_name)
            by_on_behalf_paragraph.add_run("\nDIN: " + DIN_number)
            by_on_behalf_paragraph.add_run("\n\nDate: " + date_of_meeting)
            # by_on_behalf_paragraph.add_run("\nPlace: " +place_of_meet)

            # Save the document     
            

            # Create a new document

           

            # Save the document
            # doc.save("notice_with_agenda_list.docx")

            # Provide the file for download
            filepath="notice_with_agenda_list.docx"

            company_name_under= re.sub(r'\s+', '_', company_name)

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            # AWS S3 configuration
            AWS_ACCESS_KEY = 'AKIAW2CLDMGVULENQAVM'
            AWS_SECRET_KEY = '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'
            BUCKET_NAME = 'legalnitiai'
            s3 = boto3.client('s3', aws_access_key_id=AWS_ACCESS_KEY, aws_secret_access_key=AWS_SECRET_KEY)

            file_name = f"notice_with_agenda_{company_name_under}_{formatted_time}.docx"
            filepath = f"{file_name}"
            
            heading = doc.add_heading("Notice of Meet", level=1)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = 'attachment; filename=notice_with_agenda_list.docx'
            # file_path = f"directors_table_{directors_count}_directors.docx"
            
            s3.upload_fileobj(buffer, BUCKET_NAME, filepath, ExtraArgs={
                # 'ContentType': 'application/pdf',
                'ACL': 'public-read'
            })


            print("File uploaded successfully to S3 bucket")

            # Generate a pre-signed URL for the uploaded file
            full_url = f"https://{BUCKET_NAME}.s3.amazonaws.com/{filepath}"
            print("Full URL:", full_url)

           
            return Response({'success': True, 'file_path': full_url})
