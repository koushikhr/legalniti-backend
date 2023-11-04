from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.views import View
import requests
import json
import urllib.parse
from django.http import HttpResponseBadRequest
import io
import urllib.request
import tempfile
from PIL import Image
from easyocr import Reader
from ultralytics.models.yolo import YOLO
from django.shortcuts import render
from django.http import HttpResponse
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document
from rest_framework.views import APIView
from rest_framework.response import Response
from django.views.generic import TemplateView
from docx import Document
from docx.shared import Pt
import openpyxl
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import requests
from rest_framework import status
import boto3
from botocore.exceptions import NoCredentialsError
from io import BytesIO
from datetime import datetime
import pytz
import datetime
from docx2pdf import convert
from docx.shared import Inches
import os
import re

@csrf_exempt
def noc(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        name = data.get('employeeName')
        address = data.get('officeAddress')
        company_name = data.get('companyName')

        current_datetime = datetime.datetime.now()

        # Format the current date and time
        formatted_time = current_datetime.strftime("%Y-%m-%d_%H-%M-%S")

        doc = Document()
        # Read the JSON data for Board meetings

        # Add a title above the heading
        title = doc.add_paragraph('To Whom So Ever It May Concern\n\n')
        title.style = 'Heading1'
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        intro = doc.add_paragraph()
        intro.add_run(
            f'This is to certify that, I,{name}, as the owner of the {address} hereby give consent to {company_name} for conducting their business operations at the aforementioned property. I declare that I have no objection to the company utilizing the premises for their business activities')
        intro.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        address = doc.add_paragraph()

        address.add_run(f'Address as follows\n\nC/0 {name} \n\n{address}')

        # Save the document
        # doc.save('Directors_Report.docx')
        filepath = "Directors_Report.docx"

        company_name_under = re.sub(r'\s+', '_', company_name)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # AWS S3 configuration
        AWS_ACCESS_KEY = 'AKIAW2CLDMGVULENQAVM'
        AWS_SECRET_KEY = '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'
        BUCKET_NAME = 'legalnitiai'
        s3 = boto3.client('s3', aws_access_key_id=AWS_ACCESS_KEY, aws_secret_access_key=AWS_SECRET_KEY)

        file_name = f"directors_report_{company_name_under}_{formatted_time}.docx"
        filepath = f"{file_name}"

        heading = doc.add_heading("directors_report", level=1)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=directors.docx'
        # file_path = f"directors_table_{directors_count}_directors.docx"

        s3.upload_fileobj(buffer, BUCKET_NAME, filepath, ExtraArgs={
            # 'ContentType': 'application/pdf',
            'ACL': 'public-read'
        })

        print("File uploaded successfully to S3 bucket")

        # Generate a pre-signed URL for the uploaded file
        full_url = f"https://{BUCKET_NAME}.s3.amazonaws.com/{filepath}"
        print("Full URL:", full_url)

        return JsonResponse({'success': True, 'file_path': full_url})
    else:
        return JsonResponse({"message": "Invalid request method"}, status=400)

