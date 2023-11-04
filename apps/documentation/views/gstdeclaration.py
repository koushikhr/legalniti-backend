from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.views import View
import requests
import json
import urllib.parse
from django.http import HttpResponseBadRequest
import io
import urllib.request
import tempfile
from PIL import Image
from easyocr import Reader
from ultralytics.models.yolo import YOLO
from django.shortcuts import render
from django.http import HttpResponse
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document
from rest_framework.views import APIView
from rest_framework.response import Response
from django.views.generic import TemplateView
from docx import Document
from docx.shared import Pt
import openpyxl
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import requests
from rest_framework import status
import boto3
from botocore.exceptions import NoCredentialsError
from io import BytesIO
from datetime import datetime
import pytz
import datetime
from docx2pdf import convert
from docx.shared import Inches
import os
import re

@csrf_exempt
def gstdeclaration(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        name_proprietor = data.get('name_proprietor')
        name = data.get('name')
        date = data.get('date')
        address = data.get('address')
        company_name = data.get('company_name')

        current_datetime = datetime.datetime.now()

        # Format the current date and time
        formatted_time = current_datetime.strftime("%Y-%m-%d_%H-%M-%S")

        # Create a new Document
        doc = Document()

        # Add a centered heading to the document
        heading = doc.add_heading("Declaration for Authorized Signatory", level=1)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add content to the document
        content = (
            f"{name_proprietor}\n"
            f"{address}\n"
            f"{date}\n"
            "\n"
            "Subject: Letter of Authorization for Sole Proprietor\n"
            "\n"
            "Dear GST Authority,\n"
            "\n"

        )

        # Add the content as a left-aligned paragraph
        p = doc.add_paragraph(content)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        content = (
            f"{name_proprietor}, am the owner and sole proprietor of The Sole for Soul, "
            f"located at C/0 Shilpa Kishor Gawali, {address}. "
            "I am writing this letter to authorize Vrushabh Dhanraj Chaudhari to act on my behalf in matters related to my business.\n"
            "\n"
            f"I hereby grant {name} full authority to:\n"
            "1. Represent me and my business in all official and legal matters.\n"
            "2. Enter into contracts, agreements, and transactions on behalf of the business.\n"
            "3. Make financial decisions and undertake financial transactions, including banking and payment activities.\n"
            "4. Sign documents, including but not limited to contracts, agreements, and legal forms.\n"
            "5. Access and retrieve business-related information, records, and documents.\n"
            "6. Perform any other actions necessary for the smooth operation of the business.\n"
            "\n"
            "This authorization is effective immediately and shall remain valid until revoked in writing by me. "
            f"{name} is authorized to present this letter as proof of their authority to act on my behalf.\n"
            "\n"
            f"I trust that {name} will exercise their responsibilities diligently and in the best interest of my business. "
            "I will hold them accountable for their actions and decisions made within the scope of this authorization.\n"
            "\n"
            f"Please provide {name} with any assistance they may require in carrying out their authorized duties. "
            "I kindly request that you recognize and acknowledge this letter of authorization as a valid representation of my consent "
            f"for {name} to act on behalf of my business.\n"
            "\n"
            "Thank you for your attention to this matter. Should you have any questions or require further clarification, "
            "please do not hesitate to contact me at the provided contact information."
        )

        # Add the content as a left-aligned paragraph
        p = doc.add_paragraph(content)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Add "Yours sincerely,"
        sincerely_text = "Yours sincerely,"
        sincerely_paragraph = doc.add_paragraph(sincerely_text)
        sincerely_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Add an image with signature aligned to the left bottom

        signature_text = "Signature"
        signature_paragraph = doc.add_paragraph(signature_text)
        signature_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Add additional information
        additional_info = f"{name_proprietor}\nSole Proprietor"
        additional_info_paragraph = doc.add_paragraph(additional_info)
        additional_info_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Save the document

        filepath = "declaration_document.docx"
        company_name_under = re.sub(r'\s+', '_', company_name)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # AWS S3 configuration
        AWS_ACCESS_KEY = 'AKIAW2CLDMGVULENQAVM'
        AWS_SECRET_KEY = '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'
        BUCKET_NAME = 'legalnitiai'
        s3 = boto3.client('s3', aws_access_key_id=AWS_ACCESS_KEY, aws_secret_access_key=AWS_SECRET_KEY)

        file_name = f"gst_declaration_{company_name_under}_{formatted_time}.docx"
        filepath = f"{file_name}"

        heading = doc.add_heading("declaration", level=1)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=declaration.docx'
        # file_path = f"directors_table_{directors_count}_directors.docx"

        s3.upload_fileobj(buffer, BUCKET_NAME, filepath, ExtraArgs={
            # 'ContentType': 'application/pdf',
            'ACL': 'public-read'
        })

        print("File uploaded successfully to S3 bucket")

        # Generate a pre-signed URL for the uploaded file
        full_url = f"https://{BUCKET_NAME}.s3.amazonaws.com/{filepath}"
        print("Full URL:", full_url)

        return JsonResponse({'success': True, 'file_path': full_url})

    else:
        return JsonResponse({"message": "Invalid request method"}, status=400)

# class gstdeclaration(APIView):
#
#     def post(self, request, *args, **kwargs):
#         name_proprietor = request.data.get('name_proprietor')
#         name = request.data.get('name')
#         date = request.data.get('date')
#         address = request.data.get('address')
#         company_name = request.data.get('company_name')
#
#
#         current_datetime = datetime.datetime.now()
#
#         # Format the current date and time
#         formatted_time = current_datetime.strftime("%Y-%m-%d_%H-%M-%S")
#
#
#         # Create a new Document
#         doc = Document()
#
#         # Add a centered heading to the document
#         heading = doc.add_heading("Declaration for Authorized Signatory", level=1)
#         heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#
#
#
#         # Add content to the document
#         content = (
#             f"{name_proprietor}\n"
#             f"{address}\n"
#             f"{date}\n"
#             "\n"
#             "Subject: Letter of Authorization for Sole Proprietor\n"
#             "\n"
#             "Dear GST Authority,\n"
#             "\n"
#
#         )
#
#         # Add the content as a left-aligned paragraph
#         p = doc.add_paragraph(content)
#         p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
#
#
#         content = (
#             f"{name_proprietor}, am the owner and sole proprietor of The Sole for Soul, "
#             f"located at C/0 Shilpa Kishor Gawali, {address}. "
#             "I am writing this letter to authorize Vrushabh Dhanraj Chaudhari to act on my behalf in matters related to my business.\n"
#             "\n"
#             f"I hereby grant {name} full authority to:\n"
#             "1. Represent me and my business in all official and legal matters.\n"
#             "2. Enter into contracts, agreements, and transactions on behalf of the business.\n"
#             "3. Make financial decisions and undertake financial transactions, including banking and payment activities.\n"
#             "4. Sign documents, including but not limited to contracts, agreements, and legal forms.\n"
#             "5. Access and retrieve business-related information, records, and documents.\n"
#             "6. Perform any other actions necessary for the smooth operation of the business.\n"
#             "\n"
#             "This authorization is effective immediately and shall remain valid until revoked in writing by me. "
#             f"{name} is authorized to present this letter as proof of their authority to act on my behalf.\n"
#             "\n"
#             f"I trust that {name} will exercise their responsibilities diligently and in the best interest of my business. "
#             "I will hold them accountable for their actions and decisions made within the scope of this authorization.\n"
#             "\n"
#             f"Please provide {name} with any assistance they may require in carrying out their authorized duties. "
#             "I kindly request that you recognize and acknowledge this letter of authorization as a valid representation of my consent "
#             f"for {name} to act on behalf of my business.\n"
#             "\n"
#             "Thank you for your attention to this matter. Should you have any questions or require further clarification, "
#             "please do not hesitate to contact me at the provided contact information."
#         )
#
#         # Add the content as a left-aligned paragraph
#         p = doc.add_paragraph(content)
#         p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
#
#
#
#         # Add "Yours sincerely,"
#         sincerely_text = "Yours sincerely,"
#         sincerely_paragraph = doc.add_paragraph(sincerely_text)
#         sincerely_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
#
#         # Add an image with signature aligned to the left bottom
#
#
#         signature_text="Signature"
#         signature_paragraph = doc.add_paragraph(signature_text)
#         signature_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
#
#
#         # Add additional information
#         additional_info = f"{name_proprietor}\nSole Proprietor"
#         additional_info_paragraph = doc.add_paragraph(additional_info)
#         additional_info_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
#
#         # Save the document
#
#         filepath = "declaration_document.docx"
#         company_name_under= re.sub(r'\s+', '_', company_name)
#
#         buffer = BytesIO()
#         doc.save(buffer)
#         buffer.seek(0)
#
#         # AWS S3 configuration
#         AWS_ACCESS_KEY = 'AKIAW2CLDMGVULENQAVM'
#         AWS_SECRET_KEY = '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'
#         BUCKET_NAME = 'legalnitiai'
#         s3 = boto3.client('s3', aws_access_key_id=AWS_ACCESS_KEY, aws_secret_access_key=AWS_SECRET_KEY)
#
#         file_name = f"gst_declaration_{company_name_under}_{formatted_time}.docx"
#         filepath = f"{file_name}"
#
#         heading = doc.add_heading("declaration", level=1)
#         heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#
#         response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
#         response['Content-Disposition'] = 'attachment; filename=declaration.docx'
#             # file_path = f"directors_table_{directors_count}_directors.docx"
#
#         s3.upload_fileobj(buffer, BUCKET_NAME, filepath, ExtraArgs={
#                 # 'ContentType': 'application/pdf',
#                 'ACL': 'public-read'
#             })
#
#
#         print("File uploaded successfully to S3 bucket")
#
#             # Generate a pre-signed URL for the uploaded file
#         full_url = f"https://{BUCKET_NAME}.s3.amazonaws.com/{filepath}"
#         print("Full URL:", full_url)
#
#
#         return Response({'success': True, 'file_path': full_url})



        

