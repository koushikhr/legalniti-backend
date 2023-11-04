
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.views import View
import requests
import json
import urllib.parse
from django.http import HttpResponseBadRequest
import io
import urllib.request
import tempfile
from PIL import Image
from easyocr import Reader
from ultralytics.models.yolo import YOLO
from django.shortcuts import render
from django.http import HttpResponse
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document
from rest_framework.views import APIView
from rest_framework.response import Response
from django.views.generic import TemplateView
from docx import Document
from docx.shared import Pt
import openpyxl
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import requests
from rest_framework import status
import boto3
from botocore.exceptions import NoCredentialsError
from io import BytesIO
from datetime import datetime
import pytz
import datetime
from docx2pdf import convert
from docx.shared import Inches
import os
import re



class AOA(APIView):

    def post(self, request, *args,**kwargs):

        company_name = request.data.get('companyname')
        board_meeting_numer = "GM 02/2022-23"
        date_of_meet =  request.data.get('Date')
        time_of_meet = request.data.get('Time')
        company_address = request.data.get('address')
        place_of_meet = "Bangalore"
        chairman = request.data.get('Chairman')
        end_of_meet = request.data.get('endTime')

        current_datetime = datetime.datetime.now()

        # Format the current date and time
        formatted_time = current_datetime.strftime("%Y-%m-%d_%H-%M-%S")

        doc = Document()
        def add_heading(doc, text, level=1):
            heading = doc.add_heading(text, level)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        add_heading(doc, company_name, level=1)

        # Add a new paragraph for the company address below the header
        
        doc.add_paragraph(company_address, style='BodyText')
        # Add the text to the document
        def add_content(doc, content):
            paragraph = doc.add_paragraph(content)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        meeting_minutes_text = (
            "EXTRACT OF THE MINUTES OF THE EXTRA-ORDINARY GENERAL MEETING OF THE MEMBERS\n"
            f"OF {company_name}\n"
            f"NUMBERED {board_meeting_numer} HELD ON {date_of_meet} AT {time_of_meet} \n"
            
            "SECTOR 3, HSR LAYOUT, BANGALORE-560102, KARNATAKA, INDIA THROUGH VIDEO CONFERENCING:"
        )
        # Add the text to the document
        add_content(doc, meeting_minutes_text)
        add_heading(doc, "MEMBERS AND DIRECTORS PRESENT:", level=1)
        # Create the table for director details
        directors_table = doc.add_table(rows=1, cols=5)
        directors_table.style = "Table Grid"



        # Set the column widths
        column_widths = [1, 3, 3, 3]
        for i, width in enumerate(column_widths):
            directors_table.cell(0, i).width = width

        # Add the column headers for director details
        directors_headers = directors_table.rows[0].cells
        directors_headers[0].text = "Sr. No."
        directors_headers[1].text = "Name of Directors"
        directors_headers[2].text = "Designation"
        directors_headers[3].text = "Mode of Presence"
        directors_headers[4].text = "Signature"  # Add the correct header text for the 5th column

        for i, director_data in enumerate(request.data.get('selectedDirectors', [])):
            sr_no = str(i + 1)
            name = director_data.get('NAME', '')
            designation = "Director"
            din_number = director_data.get('DIN', '')
            mode_of_presence =''
            signature=''

            row = directors_table.add_row().cells
            row[0].text = sr_no
            row[1].text = name
            row[2].text = designation
            row[3].text = mode_of_presence
            row[4].text = signature

        add_heading(doc, "INVITEES PRESENT:", level=1)

        # Create the table for invitees present
        invites_table = doc.add_table(rows=1, cols=2)
        invites_table.style = "Table Grid"

        # Set the column widths for invitees table
        column_widths_invites = [3, 3]
        for i, width in enumerate(column_widths_invites):
            invites_table.cell(0, i).width = Inches(width)

        # Add the column headers for invitees present
        invites_headers = invites_table.rows[0].cells
        invites_headers[0].text = "Name"
        invites_headers[1].text = "Mode of Presence"

        add_heading(doc, "DURATION OF THE MEETING:\n\n", level=1)
        
        add_content(doc,f"Time of Commencement of the Meeting:  {time_of_meet}")
        
        add_content(doc,f"Time of Closure of the Meeting::  {end_of_meet}")


        add_heading(doc, "ELECTION OF CHAIRMAN:", level=1)

        add_content(doc,f"Mr. {chairman}, authorised representative of SCI Investments VI, was elected as the chairman and welcomed the members to the meeting. ")


        add_heading(doc, "TAKING NOTE OF DOCUMENTS AVAILABLE FOR INSPECTION:", level=1)

        # Add the provided text
        add_content(doc, "The Chairman informed the members present that the following documents were available for the purpose of inspection by the shareholders and that the same are accessible during the continuance of the meeting:")
        add_content(doc, "1. Copy of the Notice and Explanatory Statement.")
        add_content(doc, "2. The shareholders’ agreement of the Company dated 25 May 2021, as amended by the amendment agreement I to the shareholders’ agreement, dated 03 August 2021, as further amended by the amendment agreement II to the shareholders’ agreement, dated 08th December 2022 (the “Shareholders’ Agreement”).")
        add_content(doc, "3. Draft amended Articles of Association of the Company.")
        add_content(doc, "4. All other supporting documents.")


        add_heading(doc, "ASCERTAINMENT OF QUORUM:", level=1)
        add_content(doc, "The requisite quorum for the meeting being present, the Chairman called the meeting in order. ")


        add_heading(doc, "SPECIAL BUSINESS:\n\nAGENDA ITEM NO.01\n APPROVAL FOR THE ADOPTION OF THE AMENDED AND RESTATED ARTICLES OF ASSOCIATION OF THE COMPANY:", level=1)
        # Add the provided text
        add_content(doc, "It was informed to the members present that pursuant to circular resolution passed by the board of directors on 8th December 2022 the directors have approved the amended shareholder’s agreement which was previously executed on 25th May 2021.")
        add_content(doc, "Considering the above, the Company proposes to amend and restate the Articles of Association of the Company to include various clauses of the Amended Shareholder’s Agreement.")
        add_content(doc, "The provisions of the Companies Act, 2013 require the Company to seek the approval of the shareholders for the proposed alteration of the articles of association of the Company, and accordingly, the board of directors recommends the relevant resolution for the approval of the shareholders.")
        add_content(doc, "The shareholders present took note of the same and passed the following as a Special Resolution:")
        #space for table 


        # Add the provided resolutions
        add_content(doc, "RESOLVED THAT provisions of Sections 5, 14 and other applicable provisions, if any, of the Companies Act, 2013 read with Companies (Incorporation) Rules, 2014 (including any statutory modification(s), enactment(s) or re-enactment(s) thereof for the time being in force), approval of members be and is hereby accorded for the adoption of the amended and restated Articles of Association of the Company as submitted to this meeting [duly initialled by the Chairman for the purpose of identification), in substitution, and to the entire exclusion of the existing Articles of Association of the Company.")
        add_content(doc, "RESOLVED FURTHER THAT any one of the Directors of the Company, be and is hereby severally authorized to take such steps as may be necessary for obtaining approvals, statutory, contractual or otherwise, in relation to the above and to settle all matters arising out of and incidental thereto, sign and execute all deeds, applications, documents and writings that may be required, on behalf of the Company and generally to do all acts, deeds, things, etc. as may be required to comply with all formalities in this regard, but not restricted to file the required form with the Ministry of Corporate Affairs / Registrar of Companies.")
        add_content(doc, "RESOLVED FURTHER THAT the copies of the foregoing resolutions, certified to be true by any directors, may be furnished to any person[s) as may be required.")


        # Add a heading for agendas
        add_heading(doc, "VOTE OF THANKS: ", level=1)

        add_content(doc, "With all the items of the agenda being transacted, the Rial Verma thanked the members and the Chairman for making it convenient to attend the Extra-Ordinary General Meeting and thanked them for their active participation in the meeting. Thereafter meeting declared as concluded.")
        # Add a blank paragraph for spacing
        doc.add_paragraph()

        # Add the company's name and chairman's details
        doc.add_paragraph(f"For and on behalf of the Board of Directors, {company_name}")

        # Add a blank paragraph for spacing
        doc.add_paragraph()

        # Add the line for signature
        doc.add_paragraph("_____________________________")

        # Add the chairman's details
        doc.add_paragraph(chairman)
        doc.add_paragraph("Chairman of the meeting")
        #doc.add_paragraph(f"DIN: {chairman_din}")

        # Add a blank paragraph for spacing
        doc.add_paragraph()

        # Add the date and place of the meeting
        doc.add_paragraph(f"Date: {date_of_meet}")
        doc.add_paragraph(f"Place: {place_of_meet}")

        # doc.save('Company_AOA.docx')
        filepath="company_AOA.docx"

        company_name_under= re.sub(r'\s+', '_', company_name)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

                # AWS S3 configuration
        AWS_ACCESS_KEY = 'AKIAW2CLDMGVULENQAVM'
        AWS_SECRET_KEY = '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'
        BUCKET_NAME = 'legalnitiai'
        s3 = boto3.client('s3', aws_access_key_id=AWS_ACCESS_KEY, aws_secret_access_key=AWS_SECRET_KEY)

        file_name = f"company_aoa_{company_name_under}_{formatted_time}.docx"
        filepath = f"{file_name}"
                
        heading = doc.add_heading("commpany_aoa", level=1)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=company_aoa.docx'
                # file_path = f"directors_table_{directors_count}_directors.docx"
                
        s3.upload_fileobj(buffer, BUCKET_NAME, filepath, ExtraArgs={
                    # 'ContentType': 'application/pdf',
                    'ACL': 'public-read'
                })


        print("File uploaded successfully to S3 bucket")

                # Generate a pre-signed URL for the uploaded file
        full_url = f"https://{BUCKET_NAME}.s3.amazonaws.com/{filepath}"
        print("Full URL:", full_url)

            
        return Response({'success': True, 'file_path': full_url})