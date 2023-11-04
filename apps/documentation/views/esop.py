from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.views import View
import requests
import json
import urllib.parse
from django.http import HttpResponseBadRequest
import io
import urllib.request
import tempfile
from PIL import Image
from easyocr import Reader
from ultralytics.models.yolo import YOLO
from django.shortcuts import render
from django.http import HttpResponse
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document
from rest_framework.views import APIView
from rest_framework.response import Response
from django.views.generic import TemplateView
from docx import Document
from docx.shared import Pt
import openpyxl
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import requests
from rest_framework import status
import boto3
from botocore.exceptions import NoCredentialsError
from io import BytesIO
from datetime import datetime
import pytz
import datetime
from docx2pdf import convert
from docx.shared import Inches
import os
import re


class ESOPALLOTMENT(APIView):
    
    def post(self, request, *args, **kwargs):

            company_name = request.data.get('companyname')
            board_meeting_numer = "GM 02/2022-23"
            date_of_meet =  request.data.get('Date')
            time_of_meet = request.data.get('Time')
            company_address = request.data.get('address')
            place_of_meet = "Bangalore"
            director_data = request.data.get('selectedDirectors', [])
            director = director_data.get('NAME', '')
            DIN = director_data.get('DIN', '')


            current_datetime = datetime.datetime.now()

        # Format the current date and time
            formatted_time = current_datetime.strftime("%Y-%m-%d_%H-%M-%S")
            

            doc = Document()

            def add_heading(doc, text, level=1):
                heading = doc.add_heading(text, level)
                heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            def add_content(doc, content):
                paragraph = doc.add_paragraph(content)
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
            add_heading(doc, {company_name}, level=1)

            doc.add_paragraph(company_address, style='BodyText')
            # Add the heading for the certified true copy
            doc.add_heading(f"CERTIFIED TRUE COPY OF THE RESOLUTION PASSED IN THE MEETING OF THE BOARD OF DIRECTORS OF {company_name} HELD ON {date_of_meet} AT {time_of_meet} AT THE REGISTERED OFFICE OF THE COMPANY SITUATED AT {company_address}  AT A SHORTER NOTICE:", level=1)
            # Add the heading for the agenda item
            doc.add_heading("AGENDA ITEM NO. 06:", level=1)
            doc.add_heading("TO APPROVE THE ALLOTMENT OF SHARES UNDER QUICKO ESOP 2016 TO THE OPTION HOLDERS:", level=2)

            # Add the content
            resolution_content = (
                "“RESOLVED THAT the board hereby acknowledges the receipt of exercise notice and exercise from below mentioned employees"
            )

            doc.add_paragraph(resolution_content)

            # Create a new table for employee details
            employee_table = doc.add_table(rows=1, cols=5)
            employee_table.style = "Table Grid"

            # Set the column widths
            column_widths = [1, 2, 3, 2, 2]
            for i, width in enumerate(column_widths):
                employee_table.cell(0, i).width = Inches(width)

            # Add the column headers for employee details
            employee_headers = employee_table.rows[0].cells
            employee_headers[0].text = "Sr. No."

            employee_headers[1].text = "Name of Employees"
            employee_headers[2].text = "Number of Shares Exercised"
            employee_headers[3].text = "Exercise Price Per Share (INR)"

            for i, employee_data in enumerate(request.data.get('selectedEmployee', [])):
                sr_no = str(i + 1)
                name = employee_data.get('NAME', '')
                share = employee_data.get('shares_exercised','')
                price = employee_data.get('exercised_price','')
                
                

                row = employee_table.add_row().cells
                row[0].text = sr_no
                row[1].text = name
                row[2].text = share
                row[3].text = price
                # total_shares_exercised += share 
            

            # doc.add_paragraph(f"Total Shares Exercised: {total_shares_exercised}")

            # Add the content related to ESOP resolution
            resolution_text = (
                "“RESOLVED FURTHER THAT pursuant to the provisions of Section 62(1)(b) and other applicable provisions, if any, "
                "of the Companies Act, 2013 read with Rule 12 of Companies (Share Capital and Debentures) Rules, 2014 and other "
                "rules made thereunder and the Articles of Association of the Company, the consent of the Board be and is hereby "
                "accorded to allot 144 (One Hundred Forty Four) equity shares having face value of INR 2/- (Indian Rupees Two Only) "
                "each, pursuant to the exercise of options under Quicko ESOP Plan 2016.”\n\n"
                
                "“RESOLVED FURTHER THAT the said Equity shares to be issued and allotted by the Company in the manner aforesaid "
                "shall rank pari-passu in all respects with the existing Equity Shares of the Company.”\n\n"
                
                "“RESOLVED FURTHER THAT any of the directors of the Company be and are hereby authorised to do all such acts, "
                "deeds, matters and things as may be necessary or expedient including filing of necessary documents, intimations "
                "including e-forms with regulatory authorities in connection to allotment of equity shares under Quicko ESOP 2016.”\n\n"
                
                "“RESOLVED FURTHER THAT the Share Certificates for the shares allotted as aforesaid be issued to abovementioned "
                "allottees after passing of above resolution.”"
            )

            # Add the "//CERTIFIED TRUE COPY//" heading at the center bottom of the page
            paragraph = doc.add_paragraph("//CERTIFIED TRUE COPY//")
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragraph.style = doc.styles['Heading1']


            # Add the content to the document
            add_content(doc, resolution_text)
            # Add the company's name and chairman's details
            doc.add_paragraph(f"For and on behalf of the Board of Directors, {company_name}")

            # Add a blank paragraph for spacing
            doc.add_paragraph()

            # Add the line for signature
            doc.add_paragraph("_____________________________")

            # Add the chairman's details
            doc.add_paragraph(director)
            doc.add_paragraph("Director  of the meeting")
            doc.add_paragraph(f"DIN: {DIN}")

            # Add a blank paragraph for spacing
            doc.add_paragraph()

            # Add the date and place of the meeting
            doc.add_paragraph(f"Date: {date_of_meet}")
            doc.add_paragraph(f"Place: {place_of_meet}")


            # doc.save('Company_Esop.docx')
            
            filepath="Company_Esop.docx"


            company_name_under= re.sub(r'\s+', '_', company_name)

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

                # AWS S3 configuration
            AWS_ACCESS_KEY = 'AKIAW2CLDMGVULENQAVM'
            AWS_SECRET_KEY = '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'
            BUCKET_NAME = 'legalnitiai'
            s3 = boto3.client('s3', aws_access_key_id=AWS_ACCESS_KEY, aws_secret_access_key=AWS_SECRET_KEY)

            file_name = f"esop_{company_name_under}_{formatted_time}.docx"
            filepath = f"{file_name}"
                
            heading = doc.add_heading("Attendance Sheet", level=1)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = 'attachment; filename=esop.docx'
                # file_path = f"directors_table_{directors_count}_directors.docx"
                
            s3.upload_fileobj(buffer, BUCKET_NAME, filepath, ExtraArgs={
                    # 'ContentType': 'application/pdf',
                    'ACL': 'public-read'
                })


            print("File uploaded successfully to S3 bucket")

                # Generate a pre-signed URL for the uploaded file
            full_url = f"https://{BUCKET_NAME}.s3.amazonaws.com/{filepath}"
            print("Full URL:", full_url)

            
            return Response({'success': True, 'file_path': full_url})
    

