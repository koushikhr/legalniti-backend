from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.views import View
import requests
import json
import urllib.parse
from django.http import HttpResponseBadRequest
import io
import urllib.request
import tempfile
from PIL import Image
from easyocr import Reader
from ultralytics.models.yolo import YOLO
from django.shortcuts import render
from django.http import HttpResponse
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document
from rest_framework.views import APIView
from rest_framework.response import Response
from django.views.generic import TemplateView
from docx import Document
from docx.shared import Pt
import openpyxl
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import requests
from rest_framework import status
import boto3
from botocore.exceptions import NoCredentialsError
from io import BytesIO
from datetime import datetime
import pytz
import datetime
from docx2pdf import convert
from docx.shared import Inches
import os
import re


class AMENDEDAOA(APIView):

    def post(self,request,*args,**kwargs):

            company_name = request.data.get('companyname')
            board_meeting_numer = "GM 02/2022-23"
            date_of_meet =  request.data.get('Date')
            time_of_meet = request.data.get('Time')
            company_address = request.data.get('address')
            place_of_meet = "Bangalore"
            director_data = request.data.get('selectedDirectors', [])
            if director_data:
                director = director_data[0].get('NAME', '')
                DIN = director_data[0].get('DIN', '')
            CIN = ""
            website=""
            current_datetime = datetime.datetime.now()

        # Format the current date and time
            formatted_time = current_datetime.strftime("%Y-%m-%d_%H-%M-%S")



            doc = Document()

            

            # Function to add headings
            def add_heading(doc, text, level=1):
                heading = doc.add_heading(text, level)
                heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
            def add_content(doc, content):
                paragraph = doc.add_paragraph(content, style='BodyText')
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
            add_heading(doc, company_name, level=1)


            # Add a new paragraph for the company address below the header
            
            doc.add_paragraph(company_address, style='BodyText')

            meeting_minutes_text = (
                "EXTRACT OF THE MINUTES OF THE EXTRA-ORDINARY GENERAL MEETING OF THE MEMBERS\n"
                f"OF {company_name}\n"
                f"NUMBERED {board_meeting_numer} HELD ON {date_of_meet} AT {time_of_meet} AT\n"
                
                "SECTOR 3, HSR LAYOUT, BANGALORE-560102, KARNATAKA, INDIA THROUGH VIDEO CONFERENCING:"
            )
            # Add the text to the document
            add_content(doc, meeting_minutes_text)

            # Add headings)
            add_heading(doc, "SPECIAL BUSINESS:", level=1)
            add_heading(doc, "AGENDA ITEM NO. 01:", level=1)
            add_heading(doc, "APPROVAL FOR THE ADOPTION OF THE AMENDED AND RESTATED ARTICLES OF ASSOCIATION OF THE COMPANY:", level=1)

            content = "The following resolution was passed as special resolution.\n\n" \
                    "RESOLVED THAT provisions of Sections 5, 14 and other applicable provisions, if any, " \
                    "of the Companies Act, 2013 read with Companies (Incorporation) Rules, 2014 (including any " \
                    "statutory modification(s), enactment(s) or re-enactment(s) thereof for the time being in force), " \
                    "approval of members be and is hereby accorded for the adoption of the amended and restated Articles " \
                    "of Association of the Company as submitted to this meeting (duly initialled by the Chairman for the " \
                    "purpose of identification), in substitution and to the entire exclusion of the existing Articles of " \
                    "Association of the Company.\n\n" \
                    "RESOLVED FURTHER THAT any one of the Directors of the Company, be and is hereby severally authorized " \
                    "to take such steps as may be necessary for obtaining approvals, statutory, contractual or otherwise, " \
                    "in relation to the above and to settle all matters arising out of and incidental thereto, sign and " \
                    "execute all deeds, applications, documents and writings that may be required, on behalf of the Company " \
                    "and generally to do all acts, deeds, things, etc. as may be required to comply with all formalities in " \
                    "this regard, but not restricted to file the required form with the Ministry of Corporate Affairs / " \
                    "Registrar of Companies.\n\n" \
                    "RESOLVED FURTHER THAT the copies of the foregoing resolutions, certified to be true by any directors, " \
                    "may be furnished to any person[s) as may be required.\n\n"
            doc.add_paragraph(content, style='BodyText')


            # Add the "//CERTIFIED TRUE COPY//" heading at the center bottom of the page
            paragraph = doc.add_paragraph("//CERTIFIED TRUE COPY//")
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragraph.style = doc.styles['Heading1']

            

            # Add the CIN and website with spacing
            add_content(doc, f"{CIN}\t\t\t\t\t\t{website}")

            additional_content = "For Whatfix Private Limited\n"
            add_content(doc, additional_content)

            

            director_info = (
                    "___________________________________\n"
                    f"{director}\n"
                    "Director\n"
                    f"DIN: {DIN}\n\n"
                    "Date: 17.02.2023\n"
                    "Place: Bangalore"
                )
            doc.add_paragraph(director_info, style='BodyText')

            # doc.save('amended_aoa.docx')
            filepath="amended_aoa.docx"

            company_name_under= re.sub(r'\s+', '_', company_name)

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

                    # AWS S3 configuration
            AWS_ACCESS_KEY = 'AKIAW2CLDMGVULENQAVM'
            AWS_SECRET_KEY = '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'
            BUCKET_NAME = 'legalnitiai'
            s3 = boto3.client('s3', aws_access_key_id=AWS_ACCESS_KEY, aws_secret_access_key=AWS_SECRET_KEY)

            file_name = f"amendedaoa_{company_name_under}_{formatted_time}.docx"
            filepath = f"{file_name}"
                    
            heading = doc.add_heading("amendedaoa", level=1)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = 'attachment; filename=amended_aoa.docx'
                    # file_path = f"directors_table_{directors_count}_directors.docx"
                    
            s3.upload_fileobj(buffer, BUCKET_NAME, filepath, ExtraArgs={
                        # 'ContentType': 'application/pdf',
                        'ACL': 'public-read'
                    })


            print("File uploaded successfully to S3 bucket")

                    # Generate a pre-signed URL for the uploaded file
            full_url = f"https://{BUCKET_NAME}.s3.amazonaws.com/{filepath}"
            print("Full URL:", full_url)

                
            return Response({'success': True, 'file_path': full_url})