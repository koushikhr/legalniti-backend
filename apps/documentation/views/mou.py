from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.views import View
import requests
import json
import urllib.parse
from django.http import HttpResponseBadRequest
import io
import urllib.request
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
from PIL import Image
from easyocr import Reader
from ultralytics.models.yolo import YOLO
from django.shortcuts import render
from django.http import HttpResponse
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document
from rest_framework.views import APIView
from rest_framework.response import Response
from django.views.generic import TemplateView
from docx.shared import Inches

from docx.shared import Pt
import openpyxl
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import requests
from rest_framework import status
import boto3
from botocore.exceptions import NoCredentialsError
from io import BytesIO
from datetime import datetime
import pytz
import os
import re



class mou(APIView):

    def post(self,request,*args,**kwargs):
        party1=request.data.get('party1',[])
        party2=request.data.get('party2',[])
        company_name=request.data.get('companyname','')

        #party1
        party1_name=party1.get('name','')
        party1_address=party1.get('residence','')

        #party2
        party2_name=party2.get('name','')
        party2_address=party2.get('residence','')

        current_datetime = datetime.now()

        # Format the current date and time
        formatted_time = current_datetime.strftime("%Y-%m-%d_%H-%M-%S")


        # Create a new document
        doc = Document()

        # Add the centered heading "Memorandum of Understanding"
        heading = doc.add_heading("Memorandum of Understanding", level=1)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
       


        # Add the provided content paragraphs
        content = [
            f"This Memorandum of Understanding (the “MOU”) is entered into __________(the “Effective Date”), by and between {party2_name}, with an address of {party2_address} and {party1_name}, with an address of {party1_address}, also individually referred to as “Party”, and collectively “the Parties.”",
            "WHEREAS, the Parties desire to enter into an agreement to abide by the MOU and",
            "WHEREAS, the Parties desire to memorialize certain terms and conditions of their anticipated endeavor;",
            "NOW THEREFORE, in consideration of the mutual promises and covenants contained herein, the Parties agree as follows:",
            f"1. Purpose and Scope : The purpose of this MOU is to establish a collaboration between {party2_name} and {party1_name} for the advancement of cybersecurity education and training and to provide students with real-world experience in cybersecurity by working on live projects. This collaboration will involve the following:",
            f"-{party1_name} will provide access to its comprehensive cybersecurity training platform to {party2_name} for use by its students and faculty.",
            f"- {party2_name} will promote the use of {party1_name} training platform to its students and faculty.",
            f"- {party2_name} will provide a personal trainer to {party1_name} for a period of five months to facilitate the use of its training platform.",
            f"- {party2_name} will establish an internship program for {party1_name} -certified students to gain hands-on experience in the cybersecurity industry.",
            f"- {party1['name']} will provide {party2['name']} with career guidance and support for its cybersecurity students.",
            f"- {party1_name} will provide feedback and evaluation of the nominated students' work to the University.",
            "- This MOU should not establish or create any type of formal agreement or obligation. Instead, it is an agreement between the Parties to work together in such a manner to encourage an atmosphere of collaboration and alliance in the support of an effective and efficient partnership to establish and maintain objectives and commitments with regards to all matters related to cybersecurity.",
            "2. Objectives:The Parties agrees as follows:",
            "- The Parties shall work together in a cooperative and coordinated effort so as to bring about the achievement and fulfillment of the purpose of the MOU.",
            "- This MOU does not intend to restrict the Parties to this Agreement from their involvement or participation with any other public or private individuals, agencies, or organizations.",
            "- The Parties shall mutually contribute and participate in all phases of the planning and development  of cybersecurity to the fullest extent possible.",
            "- This MOU is not intended to create any rights, benefits, and/or trust responsibilities by or between the Parties.",
            "3. Responsibilities: The responsibilities of each party are as follows:",
            f"{party1_name}:",
            "- Provide access to its comprehensive cybersecurity training platform.",
            "- Provide a personal trainer to facilitate the use of its training platform.",
            f"- Provide career guidance and support to {party2_name}'s cybersecurity students.",
            "- Assisting the Bug Bounty club in identifying vulnerabilities in our Vulnerability Disclosure Program (VDP) and Bug Bounty Program (BBP) initiatives.",
            f"- {party1_name} Private Limited will provide a private server for the university to use as their training environment, ensuring students have a secure and controlled space to practice and apply their skills.",
            f"- {party1_name} Private Limited will provide technical and learning support to the universities throughout the program.",
            f"{party2_name}:",
            f"- Promote the use of {party1_name} training platform to its students and faculty.",
            f"- Represent the collaboration between BMS College of Engineering and {party1_name}  in relevant forums and media.",
            f"- Work with {party1_name} on real-world cybersecurity projects to provide hands-on experience to students.",
            "- Provide necessary resources and support for the internship/Job program.",
            "- The University will provide a computer lab or internet access to the students to use HaxSploit's training platform virtually and attend live lectures.",
            f"4. Plan and Pricing: {party1_name} Private Limited will provide the comprehensive cybersecurity training platform and personal training to a group of 20 students selected by the university. The program will cost $500 per month and will be billed on a monthly basis.",
            "5. Intellectual Property:",
            "- All intellectual property developed by {party1['name']} or the university as part of this partnership will be jointly owned and shall be governed by a separate agreement between the parties.",
            "- Both parties agree to respect the other's intellectual property rights, including but not limited to trademarks, copyrights, patents, trade secrets, and other proprietary information.",
            "- Any intellectual property developed or created by university students, faculty, or staff in connection with the partnership will remain the property of the university, subject to the terms of any separate agreements between the parties.",
            "- Both parties agree to take reasonable steps to protect any intellectual property developed or created as part of the partnership and to comply with all applicable laws and regulations related to the protection and use of intellectual property.",
            "- In the event of any dispute related to intellectual property rights arising from this partnership, both parties will work together in good faith to resolve the dispute.",
            "6.Terms and Conditions: This MOU shall become effective on the date of signature by both parties and shall remain in effect for 5 Months.",
            "7. Termination: This Agreement may be terminated at any time by either Party upon 30 days written notice to the other Party.",
            "8. Representations and Warranties: Both Parties represent that they are fully authorized to enter into this Agreement. The performance and obligations of either Party will not violate or infringe upon the rights of any third-party or violate any other agreement between the Parties, individually, and any other person, organization, or business or any law or governmental regulation.",
            "9. Indemnity: The Parties each agree to indemnify and hold harmless the other Party, its respective affiliates, officers, agents, employees, and permitted successors and assigns against any and all claims, losses, damages, liabilities, penalties, punitive damages, expenses, reasonable legal fees and costs of any kind or amount whatsoever, which result from the negligence of or breach of this Agreement by the indemnifying party, its respective successors and assigns that occurs in connection with this Agreement. This section remains in full force and effect even after termination of the Agreement by its natural termination or the early termination by either party.",
            "10. Limitation of Liability: UNDER NO CIRCUMSTANCES SHALL EITHER PARTY BE LIABLE TO THE OTHER PARTY OR ANY THIRD PARTY FOR ANY DAMAGES RESULTING FROM ANY PART OF THIS AGREEMENT SUCH AS, BUT NOT LIMITED TO, LOSS OF REVENUE OR ANTICIPATED PROFIT OR LOST BUSINESS, COSTS OF DELAY OR FAILURE OF DELIVERY, WHICH ARE NOT RELATED TO OR THE DIRECT RESULT OF A PARTY’S NEGLIGENCE OR BREACH.",
            "11. Severability: In the event any provision of this Agreement is deemed invalid or unenforceable, in whole or in part, that part shall be severed from the remainder of the Agreement and all other provisions should continue in full force and effect as valid and enforceable.",
            "13. Waiver: The failure by either Party to exercise any right, power, or privilege under the terms of this Agreement will not be construed as a waiver of any subsequent or future exercise of that right, power, or privilege or the exercise of any other right, power, or privilege.",
            "14. Legal Fees: In the event of a dispute resulting in legal action, the successful party will be entitled to its legal fees, including, but not limited to its attorneys’ fees.",
            "15. Legal and Binding Agreement: This Agreement is legal and binding between the Parties as stated above. This Agreement may be entered into and is legal and binding in the United States and Europe. The Parties each represent that they have the authority to enter into this Agreement.",
            "16. Governing Law and Jurisdiction: The Parties agree that this Agreement shall be governed by the State and/or Country in which both Parties do business.",
            "17. Confidentiality:",
            "- Both parties agree to maintain confidentiality with respect to any confidential information disclosed during the course of this MOU, including but not limited to business plans, trade secrets, customer information, and any other proprietary information that is not otherwise publicly available.",
            "- The receiving party shall take reasonable measures to protect the confidentiality of the disclosed information, and shall not disclose or use such information for any purpose other than those related to the implementation of this MOU, without the express written consent of the disclosing party.",
            "- This obligation of confidentiality shall survive the termination of this MOU and shall remain in effect for 30 days from the date of disclosure of any confidential information.",
            "18. Entire Agreement: The Parties acknowledge and agree that this Agreement represents the entire agreement between the Parties. In the event that the Parties desire to change, add, or otherwise modify any terms, they shall do so in writing to be signed by both parties.",
            "The Parties agree to the terms and conditions set forth above as demonstrated by their signatures as follows:",
        ]

        # Add content paragraphs to the document
        for paragraph in content:
            doc.add_paragraph(paragraph)

        # Add signature lines for both parties
        doc.add_paragraph(f"{party1['name']}")
        doc.add_paragraph("Signed:")
        doc.add_paragraph("Designation:")
        doc.add_paragraph("Date:")

        doc.add_paragraph(f"{party2['name']}")
        doc.add_paragraph("Signed:")
        doc.add_paragraph("Designation:")
        doc.add_paragraph("Date:")


        
        # doc.save('amended_aoa.docx')
        filepath="Memorandum_of_Understanding.docx"

        company_name_under= re.sub(r'\s+', '_', company_name)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

                # AWS S3 configuration
        AWS_ACCESS_KEY = 'AKIAW2CLDMGVULENQAVM'
        AWS_SECRET_KEY = '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'
        BUCKET_NAME = 'legalnitiai'
        s3 = boto3.client('s3', aws_access_key_id=AWS_ACCESS_KEY, aws_secret_access_key=AWS_SECRET_KEY)

        file_name = f"Memorandum_of_Understanding_{company_name_under}_{formatted_time}.docx"
        filepath = f"{file_name}"
                    
        heading = doc.add_heading("Memorandum of Understanding", level=1)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=Memorandum_of_Understanding.docx'
                # file_path = f"directors_table_{directors_count}_directors.docx"
                    
        s3.upload_fileobj(buffer, BUCKET_NAME, filepath, ExtraArgs={
                        # 'ContentType': 'application/pdf',
                        'ACL': 'public-read'
                    })


        print("File uploaded successfully to S3 bucket")

                    # Generate a pre-signed URL for the uploaded file
        full_url = f"https://{BUCKET_NAME}.s3.amazonaws.com/{filepath}"
        print("Full URL:", full_url)

                
        return Response({'success': True, 'file_path': full_url})

